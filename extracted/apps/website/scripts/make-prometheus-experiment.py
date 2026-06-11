from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "papers" / "files" / "the-gplane-architecture-final-candidate.docx"
TARGET = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Prometheus-Experiment.docx"

shutil.copyfile(SOURCE, TARGET)
doc = Document(str(TARGET))

EXACT = {
    "THE G-PLANE ARCHITECTURE": "POWER MUST SHOW ITS WARRANT",
    "Governance Infrastructure for Autonomous Systems": "The G-Plane Architecture for Governable Autonomy",
    "Including Wards, Warrants, Authority Routing, Evidence Ledgers, Governance Kernels, and Execution Control for Autonomous Infrastructure": "Wards, Warrants, Authority Routing, Evidence Ledgers, Governance Kernels, and the Execution Constitution",
    "Aristotle Agentic Publication Edition | June 2026": "Private Prometheus Draft | June 2026",
    "This publication package was prepared with the assistance of AI tools for formatting, readability editing, publication packaging, and production workflow. The underlying thesis, substantive judgment, research direction, and final publication decisions remain the responsibility of J. D. \"Pepper\" Petersen.": "This private Prometheus draft was prepared with AI assistance for restructuring, line editing, and production workflow. The thesis, judgment, architecture, and final editorial decisions remain those of J. D. \"Pepper\" Petersen.",
    "This book begins from a practical concern that has followed autonomous systems from the field into public institutions: action is moving faster than the structures that authorize it. A system that can sense, decide, coordinate, and act at machine speed cannot be governed only by after-the-fact policy, paper controls, or human review that arrives after consequence.": "This book begins with a problem I have seen in the field, in regulated markets, and in institutions trusted with real consequence: action is moving faster than the structures that authorize it. Once a system can sense, decide, coordinate, and act at machine speed, policy that arrives after consequence is no longer governance. It is paperwork at the scene.",
    "The G-Plane is my attempt to describe a runtime architecture for that problem. The question is not whether autonomous systems can act, but how human authority remains present inside action itself: who may authorize it, where authority stops, what evidence survives, how escalation works, and how an institution can later explain what happened without pretending the machine was sovereign.": "The G-Plane is my answer to that problem. It does not ask whether machines can act. They can. It asks the harder question: how does human authority remain present inside the act itself? Who authorized it? Where does authority stop? What evidence survives? Who may revoke it? What happens when the network goes dark? How does an institution explain the act afterward without pretending the machine became sovereign?",
    "Infrastructure is civilization made operational.": "Infrastructure is where civilization stops being an idea and becomes an operating condition.",
    "Governance must become runtime architecture.": "Governance has to move into the runtime.",
    "The central claim of this architecture is simple: no consequential action should execute unless it remains admissible under a valid authority chain at the exact moment execution occurs.": "The central claim is simple enough to be dangerous: no consequential machine action should execute unless it can show valid authority at the exact moment consequence becomes real.",
    "Artificial intelligence governance discussions frequently focus on model behavior, alignment, bias, explainability, or post-hoc auditing. Those discussions are necessary, but incomplete.": "Most AI governance talk begins in the wrong place. It talks about model behavior, alignment, bias, explainability, and post-hoc audit. Those subjects are necessary. They are not enough.",
    "At that moment governance must stop being descriptive. It must become enforceable.": "At that moment governance can no longer describe the world. It has to bind it.",
    "The solution cannot therefore be additional paperwork, larger monitoring systems, or more extensive compliance reporting. The solution must be architectural.": "The answer is not more paperwork, larger dashboards, or compliance reports with better adjectives. The answer is architecture.",
    "The modern world runs on invisible systems.": "The modern world runs on systems most people never see until they fail.",
    "Governance itself becomes infrastructure.": "At that point governance itself becomes infrastructure.",
    "For me the problem first became visible not through abstract artificial intelligence theory but through practical interaction with autonomous systems operating in physical environments.": "For me, the problem did not begin in an AI seminar. It began around physical systems, real airspace, weather, telemetry, liability, and machines that did not care whether the paperwork was elegant.",
    "The naive response to this problem is often to assume that stronger monitoring solves the issue.": "The lazy answer is monitoring.",
    "It does not, monitoring observes, and governance constrains.": "It does not. Monitoring observes. Governance constrains.",
}

PRINCIPLES = {
    "The Governance Plane exists because autonomous systems operate inside environments where infrastructure consequence can occur faster than human supervisory intervention. Governance is not merely a record of what happened after the fact. It is the authority structure that must be present before irreversible consequence occurs.": "The G-Plane exists because machines can now reach consequence faster than human supervision can arrive. Governance is not the record we assemble afterward. It is the authority structure that must already be present when the act crosses into the world.",
    "Infrastructure systems may execute only when a valid authority chain remains active at execution time. That chain must remain continuous from constitutional origin to operational delegation to infrastructure consequence.": "Execution is legitimate only when the authority chain is alive at the moment of action. That chain must run from constitutional origin, through operational delegation, all the way to consequence.",
    "Governance artifacts must become executable runtime structures rather than static institutional documents. Systems operating at machine speed cannot depend on human interpretation at the moment of execution.": "Governance artifacts must become executable structures. A machine-speed system cannot wait for a human being to interpret a policy memo at the edge of execution.",
    "Distributed infrastructure often crosses operational and institutional boundaries. The architecture separates constitutional legitimacy, sovereign governance domains, infrastructure environments, delegated authority, and execution admissibility so that power does not dissolve into simple identity management.": "Distributed infrastructure crosses institutional boundaries as a matter of course. The architecture separates constitutional legitimacy, sovereign domain, local infrastructure, delegated authority, and admissibility so power does not collapse into mere login credentials.",
    "Standing permissions become dangerous in autonomous systems. Warrants provide single-use, execution-bound authority artifacts that are consumed at the execution boundary rather than persisting as broad ambient permission.": "Standing permission is dangerous when systems act continuously. Warrants turn authority into a single-use instrument, bound to the act and consumed at the boundary.",
    "Governed infrastructure must leave reconstructable proof of authority, constraint, admissibility, and consequence. The Governance Evidence Ledger preserves the institutional and technical conditions under which action occurred.": "Governed infrastructure must leave proof: authority, constraint, admissibility, and consequence. The Governance Evidence Ledger preserves the conditions under which power crossed the line.",
}

PHRASES = [
    ("The Governance Plane", "the G-Plane"),
    ("Governance Plane", "G-Plane"),
    ("autonomous infrastructure systems", "autonomous infrastructure"),
    ("infrastructure consequence", "consequence"),
    ("institutional legitimacy continuity", "legitimacy continuity"),
    ("runtime governance infrastructure", "runtime governance"),
    ("The architecture attempts to", "The architecture tries to"),
    ("This difference is central to the architecture.", "That difference is the architecture."),
    ("The distinction is practical, not decorative.", "That is not decoration. It is the whole point."),
    ("This is where the architecture becomes inspectable.", "This is where the architecture can be inspected."),
    ("The general rule is simple:", "The rule is simple:"),
]


def has_image(paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def rewrite(text):
    stripped = text.strip()
    if stripped in EXACT:
        return EXACT[stripped]
    if stripped in PRINCIPLES:
        return PRINCIPLES[stripped]
    new = text
    if len(stripped) > 180:
        for old, replacement in PHRASES:
            new = new.replace(old, replacement)
    new = new.replace(". the G-Plane", ". The G-Plane")
    new = new.replace("the G-Plane, or GPlane", "the G-Plane")
    new = new.replace("GPlane", "G-Plane")
    new = new.replace("autonomous infrastructure introduce", "autonomous infrastructure introduces")
    new = new.replace("and or", "or")
    new = new.replace("Governance Invariants→", "Governance Invariants →")
    new = new.replace("Registers→", "Registers →")
    new = new.replace("Gater→", "Gater →")
    new = new.replace("Execution→", "Execution →")
    new = new.replace("Ledger→", "Ledger →")
    new = re.sub(r"\s+([,.;:?!])", r"\1", new)
    new = re.sub(r"\s{2,}", " ", new).strip()
    return new


def add_private_note():
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Publication Notices":
            h = paragraph.insert_paragraph_before("Private Prometheus Note", style="Heading 1")
            h.paragraph_format.page_break_before = True
            note = paragraph.insert_paragraph_before(
                "This is not the publication copy. It is a private authorial experiment: sharper in voice, more direct in argument, and closer to the field experience that produced the G-Plane thesis. The architecture remains the same. The wager is style, not substance.",
                style="First Paragraph",
            )
            for run in note.runs:
                run.italic = True
            return


changed = 0
for paragraph in doc.paragraphs:
    if has_image(paragraph):
        continue
    text = paragraph.text
    if not text.strip():
        continue
    new = rewrite(text)
    if new != text:
        set_text(paragraph, new)
        changed += 1

add_private_note()

for paragraph in doc.paragraphs[:12]:
    text = paragraph.text.strip()
    if text in {"POWER MUST SHOW ITS WARRANT", "The G-Plane Architecture for Governable Autonomy", 'J. D. "Pepper" Petersen', "Private Prometheus Draft | June 2026"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text == "POWER MUST SHOW ITS WARRANT":
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(30)

doc.core_properties.title = "Power Must Show Its Warrant: The G-Plane Architecture"
doc.core_properties.author = 'J. D. "Pepper" Petersen'
doc.core_properties.comments = "Private Prometheus experiment. Not linked from site."
doc.save(str(TARGET))
print(f"changed={changed} output={TARGET}")
