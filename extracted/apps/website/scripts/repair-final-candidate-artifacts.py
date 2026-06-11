from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"

REPLACEMENTS = {
    " oes ": " does ",
    " ho ": " who ",
    " hat ": " what ",
    " nder ": " under ",
    " ossibility ": " possibility ",
    " echnically ": " technically ",
    " ay this": " may this",
    " very recursive": " every recursive",
    " ay act": " may act",
    " ay execute": " may execute",
    " ay become": " may become",
}


def repair(text: str) -> str:
    original = text
    padded = f" {text} "
    for old, new in REPLACEMENTS.items():
        padded = padded.replace(old, new)
    text = padded.strip()
    text = text.replace("Does this actor have access?.", "Does this actor have access?")
    text = text.replace("Who is acting?.", "Who is acting?")
    text = text.replace("What has been delegated?.", "What has been delegated?")
    text = text.replace("Under what constitutional structure does that authority exist?.", "Under what constitutional structure does that authority exist?")
    text = re.sub(r"\s+([,.;:?!])", r"\1", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    return text if text != original else original


doc = Document(str(DOCX))
changed = 0
for paragraph in doc.paragraphs:
    text = paragraph.text
    new = repair(text)
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)
        changed += 1

doc.save(str(DOCX))
print(f"changed={changed}")
