from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Scientific-Edition.docx"
MEDIA = ROOT / "dist" / "book" / "extracted-media"

FIGURES = [
    ("image1.png", "Figure 0.1 — Full Governance Plane Stack"),
    ("image2.png", "Figure 0.2 — Warrant Lifecycle"),
    ("image3.png", "Figure 0.3 — Sovereign Commit Boundary"),
    ("image4.png", "Figure 0.4 — Runtime Federalism Model"),
    ("image5.png", "Figure 0.5 — Revocation Propagation Model"),
    ("image6.png", "Figure 0.6 — Admissibility State Calculation"),
    ("image7.png", "Figure 0.7 — Governance Evidence Ledger Chain"),
    ("image8.png", "Figure 0.8 — Governable Machine Reference Architecture"),
]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def remove_existing_figure_plates(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    start = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Figure Plates":
            start = i
            break
    if start is None:
        return

    end = len(paragraphs)
    for i in range(start + 1, len(paragraphs)):
        style = paragraphs[i].style.name if paragraphs[i].style else ""
        if style.startswith("Heading 1") and paragraphs[i].text.strip() != "":
            end = i
            break

    for paragraph in paragraphs[start:end]:
        paragraph._element.getparent().remove(paragraph._element)


def style_caption(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(16)
    for run in paragraph.runs:
        run.italic = True
        run.font.size = Pt(9.5)


def main() -> None:
    doc = Document(str(DOCX))
    remove_existing_figure_plates(doc)

    anchor = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "THE GPLANE ARCHITECTURE":
            anchor = doc.paragraphs[i - 1]
            break
    if anchor is None:
        raise RuntimeError("Could not find manuscript start anchor.")

    anchor = insert_paragraph_after(anchor, "Figure Plates", "Heading 1")
    anchor.paragraph_format.page_break_before = True

    intro = (
        "The following plates collect the core governance diagrams used throughout this edition. "
        "They are included here as a reader’s map before the argument moves into the formal architecture."
    )
    anchor = insert_paragraph_after(anchor, intro, "First Paragraph")
    anchor.paragraph_format.space_after = Pt(12)

    for image_name, caption in FIGURES:
        image_path = MEDIA / image_name
        if not image_path.exists():
            raise FileNotFoundError(image_path)

        figure_paragraph = insert_paragraph_after(anchor)
        figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure_paragraph.paragraph_format.keep_with_next = True
        figure_paragraph.paragraph_format.space_before = Pt(10)
        run = figure_paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.35))

        caption_paragraph = insert_paragraph_after(figure_paragraph, caption)
        style_caption(caption_paragraph)
        anchor = caption_paragraph

    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
