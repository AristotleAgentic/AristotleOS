from pathlib import Path
import re
import sys

from docx import Document


target = Path(sys.argv[1])
doc = Document(str(target))

REPLACEMENTS = {
    " ecursively": " recursively",
    " igher-order": " higher-order",
    " egislation": " legislation",
    " tatute": " statute",
    " perational": " operational",
    " ivilian": " civilian",
    " arket": " market",
    " onetary": " monetary",
    " nterpretive": " interpretive",
    " xecutable": " executable",
    " onsitutional": " constitutional",
    " ontinuity": " continuity",
    " nfrastructure": " infrastructure",
    " nstitutional": " institutional",
    " econstructable": " reconstructable",
    " dministrative": " administrative",
    " ynamic": " dynamic",
    " tatic": " static",
    " istributed": " distributed",
    " ontinuous": " continuous",
    " overnance": " governance",
    " uthority": " authority",
    " xecution": " execution",
    " dmissibility": " admissibility",
    " eterministic": " deterministic",
    " containment. while": " containment while",
    " legitimacy. into": " legitimacy into",
    " governance. and": " governance and",
    " command and ivilian": " command and civilian",
    " assumption. into": " assumption into",
    " systems. rather": " systems rather",
    " infrastructure. and more": " infrastructure and more",
    " authority. and more": " authority and more",
    " hierarchy. and more": " hierarchy and more",
    " documentation. and more": " documentation and more",
    "rather than traditional governance hierarchy. The architecture behaves": "The architecture behaves",
    " infrastructure. The Governance": " infrastructure. The Governance",
    "The distinction significantly matures the architecture.": "",
    "The distinction significantly matures the architecture": "",
    "This distinction significantly matures the architecture.": "",
    "The revised Governance Plane therefore introduces a profound shift in how governance itself is understood.": "",
    "The distinction significantly shaped the architecture.": "",
    "This capability significantly matures the architecture.": "",
    "The Governance Plane therefore increasingly acts": "The Governance Plane acts",
    "The Governance Compiler therefore increasingly acts": "The Governance Compiler acts",
    "The architecture therefore increasingly introduces": "The architecture introduces",
    "therefore increasingly behaves": "therefore behaves",
    "increasingly behaves": "behaves",
    "therefore increasingly ": "therefore ",
    "increasingly increasingly": "increasingly",
}

FILLER_PATTERNS = [
    r"\bThe distinction substantially matures the architecture\.\s*",
    r"\bThe transition substantially matures the architecture\.\s*",
    r"\bThe transition significantly matures the architecture\.\s*",
    r"\bThis capability substantially matures the architecture\.\s*",
    r"\bThis layered continuity significantly strengthens runtime governance realism\.\s*",
    r"\bThis capability substantially deepens operational realism\.\s*",
    r"\bThese assumptions substantially strengthen operational realism\.\s*",
]


def polish(text):
    original = text
    for old, new in REPLACEMENTS.items():
        text = text.replace(old, new)
    for pattern in FILLER_PATTERNS:
        text = re.sub(pattern, "", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"\.\s+\.", ".", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    # Remove stranded sentence fragments caused by deleting stock editorial lines.
    text = text.replace("The .", "")
    return text if text != original else original


changed = 0
for paragraph in doc.paragraphs:
    text = paragraph.text
    new = polish(text)
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)
        changed += 1

doc.save(str(target))
print(f"changed={changed}")
