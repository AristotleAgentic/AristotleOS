from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"
MEDIA = ROOT / "dist" / "book" / "extracted-media"

PLACEMENTS = [
    ("Chapter 5", "Introducing the Governance Plane", "image1.png", "Figure 0.1 — Full Governance Plane Stack"),
    ("Chapter 14", "Warrants", "image2.png", "Figure 0.2 — Warrant Lifecycle"),
    ("Chapter 61", "The Sovereign Commit Boundary", "image3.png", "Figure 0.3 — Sovereign Commit Boundary"),
    ("Chapter 55", "Runtime Federalism and the Architecture of Shared Authority", "image4.png", "Figure 0.4 — Runtime Federalism Model"),
    ("Chapter 26", "Revocation, Narrowing, and Dynamic Authority Control", "image5.png", "Figure 0.5 — Revocation Propagation Model"),
    ("Chapter 58", "The Admissibility State", "image6.png", "Figure 0.6 — Admissibility State Calculation"),
    ("Chapter 18", "The Governance Evidence Ledger", "image7.png", "Figure 0.7 — Governance Evidence Ledger Chain"),
    ("Chapter 67", "The Governable Machine", "image8.png", "Figure 0.8 — Governable Machine Reference Architecture"),
]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def has_image(paragraph: Paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def remove_front_figure_plates(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    start = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Figure Plates":
            start = i
            break
    if start is None:
        return

    end = len(paragraphs)
    for i in range(start + 1, len(paragraphs)):
        text = paragraphs[i].text.strip()
        style = paragraphs[i].style.name if paragraphs[i].style else ""
        if text == "THE GPLANE ARCHITECTURE" and style.startswith("Heading 1"):
            end = i
            break

    for paragraph in paragraphs[start:end]:
        delete_paragraph(paragraph)


def remove_existing_inline_figures(doc: Document) -> None:
    captions = {caption for _, _, _, caption in PLACEMENTS}
    paragraphs = list(doc.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        if paragraph._element is None:
            continue
        if paragraph.text.strip() in captions:
            previous = paragraphs[i - 1] if i > 0 else None
            if previous is not None and previous._element is not None and has_image(previous):
                delete_paragraph(previous)
            delete_paragraph(paragraph)


def chapter_title_paragraphs(doc: Document) -> dict[str, Paragraph]:
    paragraphs = doc.paragraphs
    result = {}
    for i, paragraph in enumerate(paragraphs[:-1]):
        if paragraph._element is None:
            continue
        chapter = paragraph.text.strip()
        title = paragraphs[i + 1].text.strip()
        for chapter_label, title_label, _, _ in PLACEMENTS:
            if chapter == chapter_label and title == title_label:
                result[title_label] = paragraphs[i + 1]
    return result


def insertion_anchor_after_intro(doc: Document, title_paragraph: Paragraph) -> Paragraph:
    paragraphs = doc.paragraphs
    title_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._element is title_paragraph._element)
    anchor = title_paragraph
    seen_body = 0
    for paragraph in paragraphs[title_index + 1:]:
        if paragraph._element is None:
            continue
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading") and text:
            break
        if text:
            anchor = paragraph
            seen_body += 1
        if seen_body >= 2:
            break
    return anchor


def add_figure_after(anchor: Paragraph, image_name: str, caption: str) -> None:
    figure = insert_paragraph_after(anchor)
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.keep_with_next = True
    figure.paragraph_format.space_before = Pt(12)
    figure.paragraph_format.space_after = Pt(4)
    run = figure.add_run()
    run.add_picture(str(MEDIA / image_name), width=Inches(6.2))

    caption_paragraph = insert_paragraph_after(figure, caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(14)
    for run in caption_paragraph.runs:
        run.italic = True
        run.font.size = Pt(9.5)


def main() -> None:
    doc = Document(str(DOCX))
    remove_front_figure_plates(doc)
    remove_existing_inline_figures(doc)
    title_paragraphs = chapter_title_paragraphs(doc)
    missing = [title for _, title, _, _ in PLACEMENTS if title not in title_paragraphs]
    if missing:
        raise RuntimeError(f"Missing chapter anchors: {missing}")

    for _, title, image_name, caption in PLACEMENTS:
        anchor = insertion_anchor_after_intro(doc, title_paragraphs[title])
        add_figure_after(anchor, image_name, caption)

    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
