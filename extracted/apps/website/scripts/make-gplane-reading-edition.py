from pathlib import Path
import re
import shutil
import sys

from docx import Document


source = Path(sys.argv[1])
target = Path(sys.argv[2])
shutil.copyfile(source, target)

doc = Document(str(target))


HEADING_PREFIXES = ("Heading",)


def style_name(p):
    return p.style.name if p.style else "Normal"


def is_body(p):
    text = p.text.strip()
    if not text:
        return False
    style = style_name(p)
    if style.startswith(HEADING_PREFIXES):
        return False
    if text in {"Publication Notices", "AI Assistance Disclosure", "Author's Note on Scope", "Publication Contents", "Manuscript"}:
        return False
    return style in {"Normal", "Body Text", "First Paragraph"} or "Body" in style


def delete_paragraph(p):
    p._element.getparent().remove(p._element)
    p._p = p._element = None


def remove_duplicate_manuscript_front_matter():
    """The source manuscript carries an internal copyright/TOC block after the
    new publication front matter. In the reading edition that block is duplicate
    packaging, so remove it and keep the argument beginning with principles.
    """
    paragraphs = doc.paragraphs
    start = None
    end = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "Copyright Page" and style_name(p).startswith("Heading"):
            start = i
            break
    if start is None:
        return 0
    for j in range(start + 1, len(paragraphs)):
        if paragraphs[j].text.strip() == "Design Principles of the Governance Plane" and style_name(paragraphs[j]).startswith("Heading"):
            end = j
            break
    if end is None:
        return 0
    doomed = list(paragraphs[start:end])
    for p in doomed:
        delete_paragraph(p)
    return len(doomed)


def set_text(p, text):
    p.clear()
    p.add_run(text)


def split_sentences(text):
    return re.split(r"(?<=[.!?])\s+", text.strip())


def smooth_sentences(text):
    text = re.sub(r"\s+", " ", text).strip()

    replacements = [
        (r"\bThe challenge introduced by autonomy is not merely technical\. It is constitutional\.",
         "The challenge introduced by autonomy is not merely technical; it is constitutional."),
        (r"\bThe central claim of this architecture is simple\. No consequential action should",
         "The central claim of this architecture is simple: no consequential action should"),
        (r"\bThe problem is no longer merely ([^.]+)\. The problem is ([^.]+)\.",
         r"The problem is no longer merely \1; it is \2."),
        (r"\bThe question is not merely ([^.]+)\. The question is ([^.]+)\.",
         r"The question is not merely \1; it is \2."),
        (r"\bThe question is not whether ([^.]+)\. The question is how ([^.]+)\.",
         r"The question is not whether \1, but how \2."),
        (r"\bNot by ([^.]+)\. But by ([^.]+)\.",
         r"Not by \1, but by \2."),
        (r"\bNot because ([^.]+)\. But because ([^.]+)\.",
         r"Not because \1, but because \2."),
        (r"\bNot merely ([^.]+)\. Not simply ([^.]+)\. Not only ([^.]+)\. ([^.]+)\.",
         r"Not merely \1, not simply \2, and not only \3; \4."),
        (r"\bSoftware optimized\. Software monitored\. Software recommended\. Software accelerated\.",
         "Software optimized, monitored, recommended, and accelerated."),
        (r"\bThey optimize\. They adapt\. They coordinate\. They select actions",
         "They optimize, adapt, coordinate, and select actions"),
        (r"\bMarkets will reward speed\. Institutions will reward efficiency\. Emergency conditions will reward operational flexibility\.",
         "Markets will reward speed, institutions will reward efficiency, and emergency conditions will reward operational flexibility."),
        (r"\bGovernance can no longer remain external policy\. Governance must become runtime architecture\.",
         "Governance can no longer remain external policy; it must become runtime architecture."),
        (r"\bFederation is not merger\. Interoperability is not surrender\. Shared action is not shared sovereignty\.",
         "Federation is not merger, interoperability is not surrender, and shared action is not shared sovereignty."),
    ]
    for old, new in replacements:
        text = re.sub(old, new, text)

    # Replace recurring AI-ish throat-clearing with more natural argumentative turns.
    text = text.replace("This distinction is consequential because", "The difference is consequential because")
    text = text.replace("This distinction carries enormous consequence.", "The difference carries enormous consequence.")
    text = text.replace("This distinction carries profound consequence.", "The difference carries profound consequence.")
    text = text.replace("This framing carries weight.", "The framing is important.")
    text = text.replace("This realization carries consequence.", "That realization changes the architecture.")
    text = text.replace("This behavior carries consequence.", "That behavior changes the governance problem.")
    text = text.replace("This phrase carries weight.", "The phrase is doing real work.")
    text = text.replace("That phrase carries weight.", "The phrase is doing real work.")

    # Join runs of very short declarative sentences into a more varied paragraph rhythm.
    sentences = split_sentences(text)
    out = []
    i = 0
    while i < len(sentences):
        cur = sentences[i].strip()
        nxt = sentences[i + 1].strip() if i + 1 < len(sentences) else ""
        nxt2 = sentences[i + 2].strip() if i + 2 < len(sentences) else ""

        if cur and nxt and len(cur) < 56 and len(nxt) < 64 and not cur.endswith(":"):
            # "Legitimacy is not abstract. Legitimacy is operational." becomes
            # "Legitimacy is not abstract; it is operational."
            m1 = re.match(r"^([A-Z][A-Za-z ]{2,40}) is not ([^.]+)\.$", cur)
            m2 = re.match(rf"^{re.escape(m1.group(1))} is ([^.]+)\.$", nxt) if m1 else None
            if m1 and m2:
                out.append(f"{m1.group(1)} is not {m1.group(2)}; it is {m2.group(1)}.")
                i += 2
                continue

        if cur and nxt and nxt2 and all(len(s) < 48 for s in (cur, nxt, nxt2)):
            # Avoid endless three-line stacks when they are really one thought.
            out.append(f"{cur[:-1]}, {nxt[0].lower() + nxt[1:-1]}, and {nxt2[0].lower() + nxt2[1:]}")
            i += 3
            continue

        out.append(cur)
        i += 1

    text = " ".join(s for s in out if s)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def merge_adjacent_body_paragraphs():
    foreword_index = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Foreword" and style_name(p).startswith("Heading")), 0)
    i = 0
    merged = 0
    while i < len(doc.paragraphs) - 1:
        if i <= foreword_index:
            i += 1
            continue
        p = doc.paragraphs[i]
        q = doc.paragraphs[i + 1]
        if not is_body(p) or not is_body(q):
            i += 1
            continue
        a = p.text.strip()
        b = q.text.strip()
        if not a or not b:
            i += 1
            continue
        if len(a) < 520 and len(a) + len(b) < 1250:
            set_text(p, smooth_sentences(f"{a} {b}"))
            q._element.getparent().remove(q._element)
            merged += 1
            continue
        i += 1
    return merged


removed_front_matter = remove_duplicate_manuscript_front_matter()
foreword_index = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Foreword" and style_name(p).startswith("Heading")), 0)

for idx, p in enumerate(doc.paragraphs):
    if idx > foreword_index and is_body(p):
        new_text = smooth_sentences(p.text)
        if new_text != p.text:
            set_text(p, new_text)

merged = 0
for _ in range(2):
    merged += merge_adjacent_body_paragraphs()

foreword_index = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Foreword" and style_name(p).startswith("Heading")), 0)
for idx, p in enumerate(doc.paragraphs):
    if idx > foreword_index and is_body(p):
        new_text = smooth_sentences(p.text)
        if new_text != p.text:
            set_text(p, new_text)

core = doc.core_properties
core.last_modified_by = "Aristotle Agentic"
core.comments = "Reading edition pass: reduced declarative stack rhythm and merged short body paragraphs."

doc.save(str(target))
print(f"{target} removed_front_matter={removed_front_matter} merged={merged} paragraphs={len(doc.paragraphs)}")
