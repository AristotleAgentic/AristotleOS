from pathlib import Path
import re
import shutil
import sys

from docx import Document


source = Path(sys.argv[1])
target = Path(sys.argv[2])
shutil.copyfile(source, target)

doc = Document(str(target))


CORE_TERMS = {
    "meta authority envelope", "mae", "ward", "wards", "warrant", "warrants",
    "authority domain", "authority envelope", "governance invariant",
    "runtime register", "commit gate", "commit point", "physical invariant",
    "evidence ledger", "model lineage", "governance mesh", "witness",
    "admissibility", "revocation", "sovereign", "sovereignty", "federalism",
    "insurable", "insurance", "emergency", "constitutional", "implementation",
    "runtime", "execution", "telemetry", "appendix", "figure", "test case",
}

FILLER_SENTENCE_PATTERNS = [
    r"\b(The|This|That) (distinction|separation|realization|capability|requirement|transition|architecture) (?:may |can |therefore |increasingly |ultimately |substantially |significantly )?(?:becomes?|become|strengthens?|deepens?|sits|changes|defines|sharpens)[^.]*\.",
    r"\bThis is one of [^.]*\.",
    r"\bIt is one of [^.]*\.",
    r"\bThat warning runs through the architecture\.",
    r"\bThe difference changes the frame\.",
    r"\bThe separation is foundational\.",
    r"\bThe separation is critical\.",
    r"\bThe separation is especially important\.",
    r"\bThe distinction is profound\.",
    r"\bThe distinction sits [^.]*\.",
    r"\bThe realization [^.]*\.",
    r"\bThe architecture therefore [^.]*\.",
    r"\bThe Governance Plane increasingly (?:recognizes|argues|treats|assumes) that [^.]*\.",
]

LOW_VALUE_PARAGRAPH_PATTERNS = [
    r"^The distinction between .* carries enormous consequence\.",
    r"^The separation is foundational\.",
    r"^The difference is consequential because",
    r"^The transition may ultimately become one of",
    r"^This becomes one of",
    r"^That transition may ultimately become one of",
    r"^The architecture increasingly",
]


def style_name(p):
    return p.style.name if p.style else "Normal"


def is_heading(p):
    return style_name(p).startswith("Heading")


def is_front_matter_index(idx):
    # Preserve the clean publication wrapper through the manuscript marker.
    for i, p in enumerate(doc.paragraphs[:80]):
        if p.text.strip() == "Foreword" and is_heading(p):
            return idx < i
    return idx < 35


def delete_paragraph(p):
    p._element.getparent().remove(p._element)
    p._p = p._element = None


def set_text(p, text):
    p.clear()
    p.add_run(text)


def normalize(text):
    return re.sub(r"\s+", " ", text).strip()


def term_count(text):
    low = text.lower()
    return sum(1 for term in CORE_TERMS if term in low)


def is_figure_or_reference(text):
    return bool(re.match(r"^(Figure|Table|Appendix|G\.\d|[A-Z]\.\d)\b", text.strip()))


def sentence_prune(text):
    text = normalize(text)
    sentences = re.split(r"(?<=[.!?])\s+", text)
    kept = []
    for sentence in sentences:
        s = sentence.strip()
        if not s:
            continue
        protected = term_count(s) >= 2 or is_figure_or_reference(s) or len(s) > 260
        if not protected and any(re.search(pattern, s) for pattern in FILLER_SENTENCE_PATTERNS):
            continue
        kept.append(s)
    out = normalize(" ".join(kept))
    out = out.replace("The Governance Plane increasingly ", "The Governance Plane ")
    out = out.replace("increasingly increasingly", "increasingly")
    out = out.replace("therefore therefore", "therefore")
    return out


def low_value_paragraph(text):
    if is_figure_or_reference(text):
        return False
    if term_count(text) >= 3:
        return False
    if len(text) < 180 and any(re.search(pattern, text) for pattern in LOW_VALUE_PARAGRAPH_PATTERNS):
        return True
    generic_phrases = [
        "one of the defining",
        "significantly deepens the architecture",
        "substantially strengthens the architecture",
        "the architecture’s deepest warnings",
        "civilization-scale autonomous systems",
    ]
    if term_count(text) <= 1 and sum(phrase in text.lower() for phrase in generic_phrases) >= 1:
        return True
    return False


def merge_short_scientific_paragraphs():
    merged = 0
    i = 0
    while i < len(doc.paragraphs) - 1:
        p = doc.paragraphs[i]
        q = doc.paragraphs[i + 1]
        if is_heading(p) or is_heading(q):
            i += 1
            continue
        if style_name(p) not in {"Body Text", "First Paragraph", "Normal"}:
            i += 1
            continue
        if style_name(q) not in {"Body Text", "First Paragraph", "Normal"}:
            i += 1
            continue
        a = normalize(p.text)
        b = normalize(q.text)
        if not a or not b or is_figure_or_reference(a) or is_figure_or_reference(b):
            i += 1
            continue
        if len(a) + len(b) <= 1050 and (len(a) < 360 or len(b) < 360):
            set_text(p, normalize(f"{a} {b}"))
            delete_paragraph(q)
            merged += 1
            continue
        i += 1
    return merged


deleted = 0
edited = 0

for idx, p in list(enumerate(doc.paragraphs)):
    if is_heading(p) or is_front_matter_index(idx):
        continue
    text = normalize(p.text)
    if not text:
        continue
    if low_value_paragraph(text):
        delete_paragraph(p)
        deleted += 1
        continue
    new = sentence_prune(text)
    if new and new != text:
        set_text(p, new)
        edited += 1
    elif not new:
        delete_paragraph(p)
        deleted += 1

merged = 0
for _ in range(2):
    merged += merge_short_scientific_paragraphs()

core = doc.core_properties
core.title = "The G-Plane Architecture: Governance Infrastructure for Autonomous Systems"
core.author = 'J. D. "Pepper" Petersen'
core.last_modified_by = "Aristotle Agentic"
core.comments = "Scientific Publication Edition: compressed repetitive argumentation while preserving structure, core thesis, and figure references."

doc.save(str(target))
print(f"{target} deleted={deleted} edited={edited} merged={merged} paragraphs={len(doc.paragraphs)}")
