from pathlib import Path
import sys

from docx import Document


target = Path(sys.argv[1])
doc = Document(str(target))

replacements = {
    " nderstandable": " understandable",
    " eterministic": " deterministic",
    " dentify": " identify",
    " ntelligent": " intelligent",
    " ccountable": " accountable",
    " onstitutional": " constitutional",
    " olicy": " policy",
    " dmissibility": " admissibility",
    " hat actions": " what actions",
    " fter execution": " after execution",
    " he operational": " the operational",
    " ontinuously": " continuously",
    " egitimacy": " legitimacy",
    " dentity": " identity",
    " xecution": " execution",
    " uthority": " authority",
    " overnance": " governance",
    " nfrastructure": " infrastructure",
    " machinereadable": " machine-readable",
    " post hoc": " post-hoc",
    "J. D. “Pepper” Petersen Founder, Aristotle Agentic Helena, Montana": "J. D. “Pepper” Petersen. Founder, Aristotle Agentic. Helena, Montana.",
    "The question is not whether autonomous systems can act. The question is how human authority remains present inside action itself:": "The question is not whether autonomous systems can act, but how human authority remains present inside action itself:",
    "This introduces an important principle, legitimacy is not static, and legitimacy becomes continuously recomputed.": "This introduces an important principle: legitimacy is not static; it is continuously recomputed.",
    "The gate increasingly behaves as constitutional enforcement boundary.": "The gate increasingly behaves as a constitutional enforcement boundary.",
}

changed = 0
for paragraph in doc.paragraphs:
    text = paragraph.text
    new = text
    for old, repl in replacements.items():
        new = new.replace(old, repl)
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)
        changed += 1

doc.save(str(target))
print(f"changed={changed}")
