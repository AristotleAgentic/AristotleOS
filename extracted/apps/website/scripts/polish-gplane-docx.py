from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "papers" / "files" / "g-plane-runtime-warrants-technical-paper.docx"
OUT_DIR = ROOT / "dist" / "recovery"
OUT = OUT_DIR / "g-plane-runtime-warrants-technical-paper-publication-polish.docx"


INK = RGBColor(26, 29, 34)
MUTED = RGBColor(82, 91, 103)
RULE = "B7C0CC"
SHADE = "EEF2F6"
DEEP = RGBColor(31, 64, 97)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=RULE, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_keep_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if value and keep is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not value and keep is not None:
        p_pr.remove(keep)


def add_field(paragraph, field_type):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_type
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def style_header_footer(doc):
    for idx, section in enumerate(doc.sections):
        section.different_first_page_header_footer = True
        if idx == 0:
            section.first_page_header.paragraphs[0].text = ""
            section.first_page_footer.paragraphs[0].text = ""
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "The G-Plane Architecture"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        apply_para_font(hp, "Aptos", 8.0, MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = 'J. D. "Pepper" Petersen | Aristotle Agentic | Page '
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        apply_para_font(fp, "Aptos", 8.0, MUTED)
        add_field(fp, "PAGE")


def apply_run_font(run, name, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def apply_para_font(paragraph, name, size=None, color=None, bold=None, italic=None):
    for run in paragraph.runs:
        apply_run_font(run, name, size, color, bold, italic)


def first_text_index(doc, text):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return i
    return -1


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    for section in doc.sections[1:]:
        section.start_type = WD_SECTION_START.CONTINUOUS
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in (
        ("Heading 1", 15, 14, 6),
        ("Heading 2", 12.5, 10, 4),
        ("Heading 3", 11.5, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DEEP
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            paragraph.paragraph_format.space_after = Pt(0)
            continue

        if i == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(3)
            apply_para_font(paragraph, "Georgia", 22, INK, True)
            set_keep_next(paragraph)
            continue
        if i == 1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(8)
            apply_para_font(paragraph, "Georgia", 12.8, MUTED, False)
            set_keep_next(paragraph)
            continue
        if i == 2:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(18)
            apply_para_font(paragraph, "Aptos", 9.6, MUTED, False)
            continue
        if text.startswith("Abstract—"):
            paragraph.paragraph_format.left_indent = Inches(0.14)
            paragraph.paragraph_format.right_indent = Inches(0.14)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
            apply_para_font(paragraph, "Georgia", 9.5, INK)
            for run in paragraph.runs[:1]:
                run.bold = True
            continue
        if text.startswith("Index Terms—"):
            paragraph.paragraph_format.left_indent = Inches(0.14)
            paragraph.paragraph_format.right_indent = Inches(0.14)
            paragraph.paragraph_format.space_after = Pt(12)
            apply_para_font(paragraph, "Aptos", 8.8, MUTED, False)
            continue
        if text == "References":
            paragraph.style = styles["Heading 1"]
            set_keep_next(paragraph)
            continue
        if text.startswith("Appendix "):
            paragraph.style = styles["Heading 1"]
            set_keep_next(paragraph)
            if text == "Appendix A. Minimal Artifact Schemas":
                paragraph.paragraph_format.page_break_before = True
            continue
        if text in {"Executive Summary", "Contributions", "AI Assistance Disclosure"}:
            paragraph.style = styles["Heading 1"]
            set_keep_next(paragraph)
            if text == "Contributions":
                paragraph.paragraph_format.page_break_before = True
            if text == "AI Assistance Disclosure":
                paragraph.paragraph_format.page_break_before = True
            continue
        if text == "References":
            paragraph.style = styles["Heading 1"]
            set_keep_next(paragraph)
            paragraph.paragraph_format.page_break_before = True
            continue
        if text[:2].isdigit() and ". " in text[:6]:
            # Existing numbered headings are already styled in the source, but this keeps
            # any plain-text section heading from drifting.
            set_keep_next(paragraph)

    refs_at = first_text_index(doc, "References")
    if refs_at >= 0:
        for paragraph in doc.paragraphs[refs_at + 1 :]:
            if not paragraph.text.strip():
                continue
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(-0.24)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0
            apply_para_font(paragraph, "Aptos", 8.6, INK)

    disclosure_at = first_text_index(doc, "AI Assistance Disclosure")
    refs_at = first_text_index(doc, "References")
    if disclosure_at >= 0:
        stop = refs_at if refs_at > disclosure_at else len(doc.paragraphs)
        for paragraph in doc.paragraphs[disclosure_at + 1 : stop]:
            paragraph.paragraph_format.left_indent = Inches(0.14)
            paragraph.paragraph_format.right_indent = Inches(0.14)
            apply_para_font(paragraph, "Georgia", 9.5, MUTED, False)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table)
        for r, row in enumerate(table.rows):
            set_row_cant_split(row)
            if r == 0:
                set_row_repeat_header(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_margins(cell)
                if r == 0:
                    set_cell_shading(cell, SHADE)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(1.5)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        apply_run_font(run, "Aptos", 8.1, INK, bold=(True if r == 0 else None))

    style_header_footer(doc)

    doc.core_properties.title = "The G-Plane Architecture"
    doc.core_properties.subject = (
        "Runtime Warrants, Commit Gates, and Evidence for Consequential Autonomous Action"
    )
    doc.core_properties.author = 'J. D. "Pepper" Petersen'
    doc.core_properties.keywords = (
        "agentic AI, autonomous systems, runtime governance, warrants, commit gates, evidence ledgers"
    )


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    style_document(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
