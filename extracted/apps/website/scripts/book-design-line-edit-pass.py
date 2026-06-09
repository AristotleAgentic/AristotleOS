from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"

FIGURE_CAPTIONS = {
    "Figure 0.1 — Full Governance Plane Stack": "Figure 0.1 — Full Governance Plane Stack. The stack shows the constitutional chain from root authority to evidence: legitimacy begins above execution, narrows through Wards and delegated authority, becomes action-specific through Warrants, and reaches consequence only after deterministic gates and physical containment.",
    "Figure 0.2 — Warrant Lifecycle": "Figure 0.2 — Warrant Lifecycle. A Warrant is not standing permission. It is born from a proposed consequential act, evaluated against active authority and telemetry, consumed at the Commit Gate, and preserved as evidence after execution or refusal.",
    "Figure 0.3 — Sovereign Commit Boundary": "Figure 0.3 — Sovereign Commit Boundary. The boundary marks the point where institutional legitimacy must survive into runtime execution. It is where possibility becomes consequence, and where capability must yield to admissible authority.",
    "Figure 0.4 — Runtime Federalism Model": "Figure 0.4 — Runtime Federalism Model. Shared authority is represented as a runtime condition rather than a political metaphor: multiple domains may coordinate without erasing their separate legal, institutional, or sovereign authority.",
    "Figure 0.5 — Revocation Propagation Model": "Figure 0.5 — Revocation Propagation Model. Revocation must travel faster than consequence. The model shows authority narrowing from constitutional source through Wards, envelopes, warrants, gates, and local fail-closed behavior.",
    "Figure 0.6 — Admissibility State Calculation": "Figure 0.6 — Admissibility State Calculation. Admissibility is computed from live authority, telemetry, revocation state, warrant validity, domain conditions, emergency mode, synchronization, and physical invariants. The action is legitimate only if the state remains valid now.",
    "Figure 0.7 — Governance Evidence Ledger Chain": "Figure 0.7 — Governance Evidence Ledger Chain. The ledger does more than log behavior. It preserves the authority chain, runtime state, Warrant, gate decision, physical containment status, execution result, and witness evidence needed to reconstruct legitimacy after the fact.",
    "Figure 0.8 — Governable Machine Reference Architecture": "Figure 0.8 — Governable Machine Reference Architecture. The governable machine separates intelligence from authority: models may reason and agents may plan, but the kernel, Commit Gate, physical invariant gater, ledger, and witness layer decide whether power may become consequence.",
}

REPEATED_PHRASES = {
    "The Governance Plane attempts to": "The architecture attempts to",
    "The Governance Plane treats": "The architecture treats",
    "The Governance Plane behaves": "The architecture behaves",
    "The Governance Plane preserves": "The architecture preserves",
    "The Governance Plane separates": "The architecture separates",
    "The Governance Plane requires": "The architecture requires",
    "The Governance Plane assumes": "The architecture assumes",
    "The Governance Plane introduces": "The architecture introduces",
    "The Governance Plane exists to": "The architecture exists to",
    "The Governance Plane focuses": "The architecture focuses",
    "The Governance Plane does not": "The architecture does not",
}


def style_name(paragraph):
    return paragraph.style.name if paragraph.style else "Normal"


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def paragraph_has_image(paragraph: Paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


def clean_repeated_phrases(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        if paragraph._element is None or style_name(paragraph).startswith("Heading") or paragraph_has_image(paragraph):
            continue
        text = paragraph.text
        new = text
        # Leave first definition occurrences alone; vary later explanatory prose.
        if "The Governance Plane, or GPlane" not in new:
            for old, replacement in REPEATED_PHRASES.items():
                new = new.replace(old, replacement)
        new = new.replace("This difference significantly deepens the architecture.", "This difference is central to the architecture.")
        new = new.replace("The separation is foundational, the", "The separation is foundational: the")
        new = new.replace("The distinction is extremely important.", "The distinction is practical, not decorative.")
        if new != text:
            set_text(paragraph, new)
            changed += 1
    return changed


def improve_figure_captions(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in FIGURE_CAPTIONS:
            set_text(paragraph, FIGURE_CAPTIONS[text])
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(14)
            for run in paragraph.runs:
                run.italic = True
                run.font.size = Pt(9.5)
            changed += 1
    return changed


def delete_reference_architecture_appendix(doc: Document) -> int:
    paragraphs = list(doc.paragraphs)
    start = None
    end = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Reference Architecture Appendix":
            start = i
        if paragraph.text.strip() == "Appendix G — Implementation Mapping and Runtime Realization":
            end = i
            break
    if start is None or end is None or end <= start:
        return 0
    for paragraph in paragraphs[start:end]:
        delete_paragraph(paragraph)
    return end - start


def replace_section(doc: Document, heading_text: str, replacements: list[tuple[str, str]]) -> int:
    paragraphs = list(doc.paragraphs)
    heading_index = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == heading_text:
            heading_index = i
            break
    if heading_index is None:
        return 0

    end = len(paragraphs)
    for i in range(heading_index + 1, len(paragraphs)):
        if style_name(paragraphs[i]).startswith("Heading 2"):
            end = i
            break

    anchor = paragraphs[heading_index]
    for text, style in reversed(replacements):
        new_para = anchor.insert_paragraph_before(text, style=style)
        anchor._p.addnext(new_para._p)
    for paragraph in paragraphs[heading_index + 1:end]:
        if paragraph._element is not None:
            delete_paragraph(paragraph)
    return end - heading_index - 1


def tighten_closing_and_appendix(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("That continuity requires Meta Authority Envelopes:"):
            set_text(
                paragraph,
                "That continuity is carried by a small set of primitives: Meta Authority Envelopes for constitutional legitimacy; Wards for protected sovereign domains; Authority Domains for local infrastructure scope; Authority Envelopes for delegated operational power; Governance Invariants for deterministic constraints; Runtime Registers for active machine-speed state; Warrants for action-specific authority; Commit Gates for execution-boundary admissibility; Physical Invariant Gaters for hard safety enforcement; and Governance Evidence Ledgers for cryptographic institutional memory. Together they turn governance from aspiration into infrastructure.",
            )
            changed += 1

    changed += delete_reference_architecture_appendix(doc)

    changed += replace_section(
        doc,
        "G.1 Implementation Chain",
        [
            ("The implementation chain preserves the structure of the book while making each layer explicit enough for engineering, audit, procurement, insurance review, and regulatory examination.", "Body Text"),
            ("In implementation terms, the MAE functions as a signed root authority; the Ward as a protected domain namespace; the Authority Domain as an infrastructure-local enforcement scope; the Authority Envelope as a scoped delegation artifact; the Governance Invariant as a compiled deterministic constraint; the Runtime Register as active evaluable state; the Warrant as a single-use execution token; the Commit Gate as an allow/refuse/escalate boundary; the Physical Invariant Gater as a hard physical interlock; and the Governance Evidence Ledger as a hash-linked evidence chain.", "Body Text"),
        ],
    )

    changed += replace_section(
        doc,
        "G.4 Minimal Deployment Architecture",
        [
            ("A minimal deployment should contain six operational surfaces. A Governance Root Service stores MAEs, amendment rules, trust anchors, issuer authority, and root revocation state. A Ward and Domain Registry maps protected governance contexts to local infrastructure surfaces.", "Body Text"),
            ("An Envelope Issuer and Compiler issues Authority Envelopes and compiles them into deterministic Governance Invariants. A Runtime Governance Kernel loads active invariants into Runtime Registers, verifies telemetry and revocation state, and prepares warrant eligibility decisions.", "Body Text"),
            ("A Warrant Service and Commit Gate issue single-use Warrants and enforce allow, refuse, narrow, escalate, or suspend decisions immediately before consequence. An Evidence Ledger and Witness Layer records the authority chain, register state, Warrant issuance, commit decision, physical gating, execution result, and reconciliation evidence.", "Body Text"),
            ("A stronger deployment adds hardware attestation, Model Lineage Certificates, threshold witness signatures, external ledger anchoring, domain-specific adapters, formal verification of invariant compilation, and disaster recovery procedures for revocation under degraded connectivity.", "Body Text"),
        ],
    )

    changed += replace_section(
        doc,
        "G.5 Commit Decision Contract",
        [
            ("Every governed action should produce a commit decision contract. The contract need not be exposed to end users, but it must exist as a machine-verifiable record that can be reconstructed later.", "Body Text"),
            ("At minimum, the contract should identify the proposed action, hash the action and material parameters, bind the action to an MAE, Ward, Authority Domain, Authority Envelope, invariant set, Runtime Register snapshot, Warrant, revocation vector, Commit Gate decision, reason code, physical-gate status, and resulting Governance Evidence Ledger record.", "Body Text"),
            ("This is where the architecture becomes inspectable. A regulator, insurer, court, operator, or internal review body should be able to determine which authority chain existed, which constraints were active, which telemetry was relied upon, which Warrant was consumed, why the gate allowed or refused the action, and how the consequence was recorded.", "Body Text"),
        ],
    )

    changed += replace_section(
        doc,
        "G.7 Implementation Test Cases",
        [
            ("A credible implementation should pass a small set of structural tests before it is treated as a governed autonomous system.", "Body Text"),
            ("Root Authority Test: no Ward, Authority Domain, Authority Envelope, or Warrant can be created without a valid MAE chain. Ward Binding Test: no delegated authority can execute unless it is bound to a protected governance context. Domain Locality Test: authority valid in one domain cannot silently execute in another without reconciliation.", "Body Text"),
            ("Invariant Compilation Test: every enforceable runtime constraint can be traced back to its source authority artifact. Register Freshness Test: Commit Gate evaluation fails, narrows, or escalates when required runtime state is stale. Single-Use Warrant Test: a Warrant cannot be replayed, reused, broadened, transferred, or executed outside its window.", "Body Text"),
            ("Commit Boundary Test: no consequential action reaches infrastructure without a recorded gate decision. Physical Containment Test: physical constraints override software authority when hard safety limits are violated. Ledger Reconstruction Test: an external reviewer can reconstruct the authority chain, runtime state, Warrant, gate decision, and execution result after the fact.", "Body Text"),
        ],
    )

    return changed


def format_front_matter(doc: Document) -> None:
    for i, paragraph in enumerate(doc.paragraphs[:20]):
        text = paragraph.text.strip()
        if text == "THE G-PLANE ARCHITECTURE":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(28)
                run.bold = True
        elif text in {"Governance Infrastructure for Autonomous Systems", 'J. D. "Pepper" Petersen', "Aristotle Agentic Publication Edition | June 2026"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text == "Publication Notices":
            paragraph.paragraph_format.page_break_before = True


doc = Document(str(DOCX))
changed = 0
changed += improve_figure_captions(doc)
changed += clean_repeated_phrases(doc)
changed += tighten_closing_and_appendix(doc)
format_front_matter(doc)

doc.core_properties.comments = "Final candidate with inline figures, improved captions, tightened appendix, and book-design line edit."
doc.save(str(DOCX))
print(f"changed={changed}")
