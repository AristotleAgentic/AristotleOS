from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"


REWRITES = {
    "Computational Rule of Law": [
        "The phrase computational rule of law should be read narrowly. It does not mean courts replaced by code, justice reduced to software, or constitutional judgment automated away. It means that systems capable of producing consequence must be unable to bypass the legitimacy conditions that make power lawful, reviewable, and bounded.",
        "Rule of law has never meant that rules merely exist. Arbitrary systems can have rules. Rule of law means power is constrained by authority that can be understood, challenged, reviewed, and applied without hidden exception. Autonomous systems threaten that settlement because execution can outrun review, delegation can become opaque, and evidence can fragment across platforms.",
        "The Governance Plane answers by moving lawful process to the commit boundary. Before consequence occurs, the system must show authority, protected context, delegated scope, current state, revocation status, physical limits, and evidence obligations. The Warrant is the runtime analog of lawful process: not a symbol of trust, but a proof that this act may proceed under this authority now.",
        "Computational rule of law is therefore the doctrine behind the mechanism. Compliance asks whether a system later appeared to follow a rule. Computational rule of law asks whether illegitimate consequence was structurally unreachable before action. That is the stronger standard autonomous infrastructure will require.",
    ],
    "Admissibility and the Physics of Legitimate Action": [
        "Admissibility is the bridge between authority and action. It is not a general property of a system. It is a present-tense judgment about a proposed consequence: may this action cross into the world under the authority, state, constraints, and evidence available now?",
        "The physics language is deliberate. Consequence has direction, timing, and irreversibility. Before execution, an action is possibility. After execution, it becomes world-state mutation: money moves, access changes, vehicles move, systems isolate, power routes, records alter, or physical systems actuate. Governance has force only if it binds before that transition.",
        "A system may possess a credential and still be inadmissible. A model may be accurate and still lack authority. A plan may be efficient and still violate a Ward. A Warrant may exist but become unusable when telemetry degrades or revocation arrives. Admissibility is where these conditions are resolved into a yes or no.",
        "This chapter establishes the concept. Later chapters describe the state machinery that computes it and the gates that enforce it. The distinction is important: admissibility is the judgment, the admissibility state is the live condition set, and the Commit Gate is the mechanism that refuses or allows execution.",
    ],
    "The Admissibility State": [
        "The admissibility state is the live condition set behind the judgment. It is not a label attached to the system. It is the current arrangement of authority, telemetry, revocation, Warrant validity, Ward context, domain conditions, emergency posture, synchronization, and physical invariants.",
        "Traditional authorization treats permission as relatively stable. Autonomous systems make that unsafe. A valid actor can become inadmissible because the environment changed. A legitimate mission can narrow because emergency authority expired. A route can become invalid because the system crossed into another Ward. The admissibility state captures those changes while the system is operating.",
        "Identity answers who is acting. Authority answers what has been delegated. Legitimacy answers where that authority comes from. Admissibility answers whether the proposed consequence may occur now. The state exists so the Commit Gate can answer that final question without pretending old permission is still enough.",
    ],
    "The Sovereign Commit Boundary": [
        "The sovereign commit boundary is the institutional side of the execution line. It asks not merely whether an action is technically ready, but whether sovereign or institutional authority has survived to the point where consequence becomes irreversible.",
        "Every domain has such a boundary. Financial systems have settlement boundaries. Military systems have weapons-release boundaries. Public agencies have decision boundaries. Industrial systems have actuation boundaries. Autonomous systems add speed and opacity, but they do not eliminate the need for a legitimate crossing.",
        "This is different from the Commit Gate. The Gate is the enforcement mechanism. The sovereign commit boundary is the constitutional place the Gate protects. It marks the point where capability must yield to authority and where permission must become admissibility.",
        "If the boundary is weak, governance becomes theatrical: law exists, policy exists, dashboards exist, but power crosses before legitimacy binds. If the boundary is strong, the system may be fast without becoming sovereign. That is the operational meaning of human authority in machine-speed infrastructure.",
    ],
    "The Last Boundary": [
        "The last boundary is the book's final image of the same problem. Before the boundary, a system may reason, simulate, draft, plan, route, recommend, or prepare. After the boundary, the world has changed.",
        "The point of the architecture is to make that crossing governable. The Meta Authority Envelope supplies root legitimacy. The Ward supplies protected context. The Authority Domain locates consequence. The Authority Envelope scopes delegation. The Runtime Register supplies current state. The Warrant carries action-specific authority. The Commit Gate decides. Physical Invariant Gaters hold the hard limits. The Evidence Ledger remembers.",
        "This boundary is where many AI governance discussions stop too early. Model behavior, explanations, safety claims, and compliance policies matter, but they do not replace the moment of consequence. An accurate recommendation can still be unauthorized. A safe plan can still be outside its Ward. A useful action can still be inadmissible.",
        "The last boundary therefore gives the reader a practical test: when this system is about to change the world, what proves that it may do so? If the answer cannot be shown, the action should not cross.",
    ],
    "The Execution Constitution": [
        "The Execution Constitution is the runtime settlement among humans, institutions, machines, and consequence. It states the central rule of the architecture: humans and human institutions remain the source of legitimacy; machines may execute only within bounded, admissible delegation.",
        "This settlement separates four things autonomous systems tend to blur. Cognition is the system's capacity to infer, plan, recommend, and optimize. Authority is the legitimate power to permit or forbid consequence. Execution is the operational transition from internal state to external change. Consequence is the changed world that follows. The architecture exists because these cannot safely collapse into one another.",
        "A model may be brilliant without possessing authority. An agent may coordinate beautifully without being entitled to expand its jurisdiction. A platform may optimize efficiently without being allowed to weaken escalation, revocation, evidence, or physical limits. Capability may support execution, but it cannot become the source of rightful power.",
        "The Execution Constitution is therefore not another name for the whole book. It is the settlement the book defends: cognition may propose, authority must authorize, the Commit Gate must admit, and evidence must preserve the chain. That is how human constitutional authority survives machine-speed action.",
    ],
}


READER_MAP = [
    ("Heading 1", "How to Read This Architecture"),
    (
        "First Paragraph",
        "The Governance Plane is easiest to read as a chain from legitimacy to consequence. The book begins with the control problem, then builds the primitives that carry authority toward execution, then tests those primitives under failure, federation, emergency, and physical consequence.",
    ),
    (
        "Body Text",
        "The primitives have distinct jobs. Meta Authority Envelopes establish who may create authority. Wards define the protected context on whose behalf authority exists. Authority Domains locate the infrastructure environment where consequence will occur. Authority Envelopes delegate bounded operational scope. Governance Invariants make constraints executable. Runtime Registers hold live state. Warrants carry action-specific authority. Commit Gates admit or refuse execution. Physical Invariant Gaters keep hard consequences unreachable. Evidence Ledgers preserve the chain after action.",
    ),
    (
        "Body Text",
        "Several boundary terms recur, but they are not synonyms. Admissibility is the judgment that an act may proceed now. The admissibility state is the live condition set used to make that judgment. The Commit Gate is the enforcement mechanism. The sovereign commit boundary is the constitutional place that must be protected. The Last Boundary is the final image of the same crossing: possibility becoming consequence.",
    ),
    (
        "Body Text",
        "Read the architecture in that order and the book becomes less a list of concepts than a single motion: authority begins above the machine, narrows as it approaches action, proves itself at the boundary, and leaves evidence behind.",
    ),
]


def has_image(paragraph: Paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def set_text(paragraph: Paragraph, text: str) -> None:
    ppr = paragraph._p.pPr
    for child in list(paragraph._p):
        if ppr is not None and child is ppr:
            continue
        paragraph._p.remove(child)
    paragraph.add_run(text)


def chapter_bounds(doc: Document, title: str, occurrence: int = 1) -> tuple[int, int] | None:
    seen = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == title:
            seen += 1
            if seen == occurrence:
                end = len(doc.paragraphs)
                for j in range(i + 1, len(doc.paragraphs)):
                    if doc.paragraphs[j].style.name == "Heading 1":
                        end = j
                        break
                return i, end
    return None


def replace_chapter(doc: Document, title: str, paragraphs: list[str], occurrence: int = 1) -> bool:
    bounds = chapter_bounds(doc, title, occurrence)
    if not bounds:
        return False
    start, end = bounds
    heading = doc.paragraphs[start]
    images = []
    body = []
    for paragraph in list(doc.paragraphs[start + 1 : end]):
        if has_image(paragraph) or paragraph.text.strip().startswith("Figure 0."):
            images.append(paragraph)
        elif paragraph.text.strip():
            body.append(paragraph)

    anchor = heading
    inserted = []
    for idx, text in enumerate(paragraphs):
        anchor = insert_after(anchor, text, "First Paragraph" if idx == 0 else "Body Text")
        inserted.append(anchor)

    if images:
        figure_anchor = inserted[min(1, len(inserted) - 1)]
        for image_para in images:
            image_para._p.getparent().remove(image_para._p)
            figure_anchor._p.addnext(image_para._p)
            figure_anchor = image_para

    for paragraph in body:
        if paragraph._element is not None:
            delete_paragraph(paragraph)
    return True


def insert_reader_map(doc: Document) -> bool:
    if any(p.text.strip() == "How to Read This Architecture" for p in doc.paragraphs):
        return False
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "Design Principles of the Governance Plane":
            anchor = paragraph
            break
    if anchor is None:
        return False
    for style, text in reversed(READER_MAP):
        new_para = OxmlElement("w:p")
        anchor._p.addprevious(new_para)
        paragraph = Paragraph(new_para, anchor._parent)
        paragraph.add_run(text)
        paragraph.style = style
    return True


def demote_microchapter(doc: Document, title: str, new_title: str, occurrence: int = 1) -> bool:
    bounds = chapter_bounds(doc, title, occurrence)
    if not bounds:
        return False
    start, _ = bounds
    heading = doc.paragraphs[start]
    set_text(heading, new_title)
    heading.style = "Heading 2"
    # Delete the preceding "Chapter N" marker when it belongs to this demoted item.
    prev = doc.paragraphs[start - 1] if start > 0 else None
    if prev is not None and prev.style.name == "Heading 2" and prev.text.strip().startswith("Chapter "):
        delete_paragraph(prev)
    return True


def clean_text_artifacts(doc: Document) -> int:
    replacements = {
        "he administrative state.": "the administrative state",
        "he runtime state.": "the runtime state",
        "nvisible sovereignty": "invisible sovereignty",
        "nstitutions. To infrastructure.": "institutions to infrastructure.",
        "ssingular": "singular",
        "id a person technically approve this system?": "did a person technically approve this system?",
        "irect supervision.": "direct supervision",
        "oftware authorization.": "software authorization",
        "hysics-bound admissibility enforcement": "physics-bound admissibility enforcement",
        "hysics-bound consequence governance": "physics-bound consequence governance",
        "nformational execution.": "informational execution",
        "hysical consequence reachability": "physical consequence reachability",
    }
    changed = 0
    for paragraph in doc.paragraphs:
        if paragraph._element is None or has_image(paragraph):
            continue
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_text(paragraph, new_text)
            changed += 1
    return changed


def main() -> None:
    doc = Document(str(DOCX))
    inserted = insert_reader_map(doc)
    rewritten = 0
    for title, paragraphs in REWRITES.items():
        if title == "The Execution Constitution":
            # Preserve the short practical note and rewrite the long chapter.
            rewritten += int(replace_chapter(doc, title, paragraphs, occurrence=2))
        else:
            rewritten += int(replace_chapter(doc, title, paragraphs))

    demoted = 0
    demoted += int(demote_microchapter(doc, "The Execution Constitution", "Execution Constitution: Practical Test", occurrence=1))
    demoted += int(demote_microchapter(doc, "The Claims of the Architecture", "The Five Claims", occurrence=1))
    artifacts = clean_text_artifacts(doc)
    doc.core_properties.comments = "Final structural editorial pass: added reader map, merged microchapters, clarified boundary concepts."
    doc.save(str(DOCX))
    print(f"reader_map_inserted={inserted} chapters_rewritten={rewritten} microchapters_demoted={demoted} artifacts_cleaned={artifacts} output={DOCX}")


if __name__ == "__main__":
    main()
