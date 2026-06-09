from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Pepper-Experiment.docx"

REPLACEMENTS = {
    "the G-Plane, or GPlane": "the G-Plane",
    "GPlane": "G-Plane",
    ". the G-Plane": ". The G-Plane",
    "Most governance theories emerge inside institutional or academic environments. the G-Plane": "Most governance theories emerge inside institutional or academic environments. The G-Plane",
    "autonomous infrastructure introduce": "autonomous infrastructure introduces",
    "and or": "or",
    "Governance Invariants→": "Governance Invariants →",
    "Registers→": "Registers →",
    "Gater→": "Gater →",
    "Execution→": "Execution →",
    "Ledger→": "Ledger →",
    "Envelope→": "Envelope →",
    "Domain→": "Domain →",
    "Ward→": "Ward →",
    "Warrant→": "Warrant →",
    "Gate→": "Gate →",
    "The G-Plane Architecture: The G-Plane Architecture for Governable Autonomy": "Power Must Show Its Warrant: The G-Plane Architecture for Governable Autonomy",
    "this exact consequential act remain admissible now": "this exact consequential act remain admissible now",
}


doc = Document(str(DOCX))
changed = 0
for paragraph in doc.paragraphs:
    text = paragraph.text
    new = text
    for old, replacement in REPLACEMENTS.items():
        new = new.replace(old, replacement)
    new = re.sub(r"\s+([,.;:?!])", r"\1", new)
    new = re.sub(r"\s{2,}", " ", new).strip()
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)
        changed += 1

doc.save(str(DOCX))
print(f"changed={changed}")
