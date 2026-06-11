from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


target = Path(sys.argv[1])
doc = Document(str(target))


INK = RGBColor(23, 21, 18)
GREEN = RGBColor(23, 53, 43)
MID_GREEN = RGBColor(49, 89, 71)
MUTED = RGBColor(82, 103, 92)


def set_style_font(style, *, name="Georgia", size=11.2, color=INK, bold=None, italic=None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.color.rgb = color
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic


def configure_style(style_name, *, size, color=INK, bold=False, italic=False, before=0, after=8, line=1.52, align=None, keep=False):
    if style_name not in doc.styles:
        return
    style = doc.styles[style_name]
    set_style_font(style, size=size, color=color, bold=bold, italic=italic)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep
    if align is not None:
        pf.alignment = align


for section in doc.sections:
    section.top_margin = Inches(0.78)
    section.right_margin = Inches(0.86)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.86)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

configure_style("Normal", size=11.2, after=8.6, line=1.52)
configure_style("Body Text", size=11.2, after=8.6, line=1.52)
configure_style("First Paragraph", size=11.6, after=8.6, line=1.52)
configure_style("Heading 1", size=20, color=GREEN, bold=True, before=30, after=11.5, line=1.15, keep=True)
configure_style("Heading 2", size=15, color=MID_GREEN, bold=True, before=20, after=7.2, line=1.2, keep=True)
configure_style("Heading 3", size=12.5, color=MUTED, bold=True, before=14, after=6, line=1.2, keep=True)

for style in doc.styles:
    if style.name.startswith("Compact"):
        set_style_font(style, size=10.8, color=INK)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.13)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.35


def paragraph_text(paragraph):
    return paragraph.text.strip()


def set_run_defaults(paragraph, *, name="Georgia", size=None, color=None, bold=None, italic=None):
    for run in paragraph.runs:
        run.font.name = name
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic


for paragraph in doc.paragraphs:
    text = paragraph_text(paragraph)
    if not text:
        continue
    style_name = paragraph.style.name if paragraph.style else "Normal"

    # The generated PDF treats these as cover/title furniture. Make the Word
    # document carry the same visual intent directly.
    if text == "THE G-PLANE ARCHITECTURE":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(158)
        paragraph.paragraph_format.space_after = Pt(13)
        set_run_defaults(paragraph, size=34, color=GREEN, bold=True)
    elif text == "Governance Infrastructure for Autonomous Systems":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(13)
        set_run_defaults(paragraph, size=15, color=MUTED, italic=True)
    elif text.startswith("Including Wards, Warrants"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(13)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run_defaults(paragraph, size=12, color=MUTED, italic=True)
    elif text.startswith('J. D. "Pepper" Petersen') or text.startswith("J. D. “Pepper” Petersen"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(42)
        set_run_defaults(paragraph, size=13, color=INK)
    elif text == "Aristotle Agentic" or "Publication Edition" in text:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(12)
        set_run_defaults(paragraph, size=12, color=MUTED, italic=("Publication Edition" in text))
    elif text in {"Publication Notices", "Author's Note on Scope", "Publication Contents", "Manuscript"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(30)
        paragraph.paragraph_format.space_after = Pt(11.5)
        paragraph.paragraph_format.keep_with_next = True
        set_run_defaults(paragraph, size=20, color=GREEN, bold=True)
    elif style_name.startswith("Heading"):
        # Preserve heading levels but eliminate residue from earlier manuscript
        # formatting and align them with the PDF hierarchy.
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if style_name == "Heading 1":
            set_run_defaults(paragraph, size=20, color=GREEN, bold=True)
        elif style_name == "Heading 2":
            set_run_defaults(paragraph, size=15, color=MID_GREEN, bold=True)
        elif style_name == "Heading 3":
            set_run_defaults(paragraph, size=12.5, color=MUTED, bold=True)
    elif style_name.startswith("Compact"):
        if not re.match(r"^[•\-\u2022]\s+", text):
            paragraph.clear()
            run = paragraph.add_run(f"• {text}")
            run.font.name = "Georgia"
            run.font.size = Pt(10.8)
            run.font.color.rgb = INK
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.13)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.35
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(8.6)
        paragraph.paragraph_format.line_spacing = 1.52
        if style_name == "First Paragraph":
            set_run_defaults(paragraph, size=11.6, color=INK)
        else:
            set_run_defaults(paragraph, size=11.2, color=INK)

core = doc.core_properties
core.title = "The G-Plane Architecture: Governance Infrastructure for Autonomous Systems"
core.author = 'J. D. "Pepper" Petersen'
core.last_modified_by = "Aristotle Agentic"
core.comments = "Word document formatted to match the clean PDF book style."

doc.save(str(target))
print(target)
