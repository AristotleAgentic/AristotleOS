from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt, RGBColor


target = Path(sys.argv[1])
doc = Document(str(target))

disclaimer_heading = "AI Assistance Disclosure"
disclaimer_text = (
    "This publication package was prepared with the assistance of AI tools for formatting, "
    "readability editing, publication packaging, and production workflow. The underlying thesis, "
    "substantive judgment, research direction, and final publication decisions remain the "
    "responsibility of J. D. \"Pepper\" Petersen."
)


def set_para_style(paragraph, *, heading=False):
    paragraph.paragraph_format.space_after = Pt(8.6)
    paragraph.paragraph_format.line_spacing = 1.52
    for run in paragraph.runs:
        run.font.name = "Georgia"
        run.font.size = Pt(20 if heading else 11.2)
        run.font.color.rgb = RGBColor(23, 53, 43) if heading else RGBColor(23, 21, 18)
        run.bold = bool(heading)


if not any(p.text.strip() == disclaimer_heading for p in doc.paragraphs):
    insert_after = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("The work is provided for research"):
            insert_after = i
            break
    if insert_after is None:
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == "Publication Notices":
                insert_after = i
                break

    if insert_after is None:
        raise SystemExit("Could not find publication notices insertion point")

    base = doc.paragraphs[insert_after]
    heading = base.insert_paragraph_before(disclaimer_heading)
    body = heading.insert_paragraph_before(disclaimer_text)

    # insert_paragraph_before inserts above the anchor, so move the new nodes after
    # the selected paragraph while preserving Word structure.
    base_el = base._element
    heading_el = heading._element
    body_el = body._element
    heading_el.getparent().remove(heading_el)
    body_el.getparent().remove(body_el)
    base_el.addnext(body_el)
    base_el.addnext(heading_el)

    set_para_style(heading, heading=True)
    set_para_style(body)

core = doc.core_properties
core.last_modified_by = "Aristotle Agentic"
core.comments = "Book package includes AI assistance disclosure in publication notices."

doc.save(str(target))
print(target)
