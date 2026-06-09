from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Prometheus-Experiment.docx"

REPLACEMENTS = {
    "The G-Plane, or G-Plane,": "The G-Plane",
    "the G-Plane, or G-Plane,": "the G-Plane",
    "Governance must become runtime architecture.": "Governance has to move into the runtime.",
    "The naive response to this problem is often to assume that stronger monitoring solves the issue. It does not, monitoring observes, and governance constrains.": "The lazy answer is monitoring. It does not. Monitoring observes. Governance constrains.",
    "Most AI governance talk begins in the wrong place.": "Most AI governance talk begins in the wrong place.",
    "The answer is not more paperwork, larger dashboards, or compliance reports with better adjectives": "The answer is not more paperwork, larger dashboards, or compliance reports with better adjectives",
    "Consequence without admissibility": "Consequence without admissibility",
    "The central claim of this architecture is simple:": "The central claim is simple enough to be dangerous:",
    "This private Prometheus draft was prepared": "This private Prometheus draft was prepared",
}

TARGET_SENTENCES = {
    "The modern world runs on invisible systems.": "The modern world runs on systems most people never see until they fail.",
    "Governance itself becomes infrastructure.": "At that point governance itself becomes infrastructure.",
    "For me the problem first became visible not through abstract artificial intelligence theory but through practical interaction with autonomous systems operating in physical environments.": "For me, the problem did not begin in an AI seminar. It began around physical systems, real airspace, weather, telemetry, liability, and machines that did not care whether the paperwork was elegant.",
    "Infrastructure is civilization made operational.": "Infrastructure is where civilization stops being an idea and becomes an operating condition.",
    "The solution cannot therefore be additional paperwork, larger monitoring systems, or more extensive compliance reporting. The solution must be architectural.": "The answer is not more paperwork, larger dashboards, or compliance reports with better adjectives. The answer is architecture.",
    "Artificial intelligence governance discussions frequently focus on model behavior, alignment, bias, explainability, or post-hoc auditing. Those discussions are necessary, but incomplete.": "Most AI governance talk begins in the wrong place. It talks about model behavior, alignment, bias, explainability, and post-hoc audit. Those subjects are necessary. They are not enough.",
}


def has_image(paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


doc = Document(str(DOCX))
changed = 0
for paragraph in doc.paragraphs:
    if has_image(paragraph):
        continue
    text = paragraph.text
    new = text
    stripped = text.strip()
    if stripped in TARGET_SENTENCES:
        new = TARGET_SENTENCES[stripped]
    else:
        for old, replacement in REPLACEMENTS.items():
            new = new.replace(old, replacement)
        for old, replacement in TARGET_SENTENCES.items():
            new = new.replace(old, replacement)
    new = new.replace("Governance Invariants→", "Governance Invariants →")
    new = new.replace("Registers→", "Registers →")
    new = new.replace("Gater→", "Gater →")
    new = new.replace("Execution→", "Execution →")
    new = new.replace("Ledger→", "Ledger →")
    new = re.sub(r"\s+([,.;:?!])", r"\1", new)
    new = re.sub(r"\s{2,}", " ", new).strip()
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)
        changed += 1

doc.save(str(DOCX))
print(f"changed={changed}")
