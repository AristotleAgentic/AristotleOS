from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Pepper-Experiment.docx"


TARGETED_REWRITES = {
    "THE G-PLANE ARCHITECTURE": "POWER MUST SHOW ITS WARRANT",
    "Governance Infrastructure for Autonomous Systems": "The G-Plane Architecture for Governable Autonomy",
    "Including Wards, Warrants, Authority Routing, Evidence Ledgers, Governance Kernels, and Execution Control for Autonomous Infrastructure": "Wards, Warrants, Authority Routing, Evidence Ledgers, Governance Kernels, and the Execution Constitution",
    "Aristotle Agentic Publication Edition | June 2026": "Private Experimental Author's Draft | June 2026",
    "This publication package was prepared with the assistance of AI tools for formatting, readability editing, publication packaging, and production workflow. The underlying thesis, substantive judgment, research direction, and final publication decisions remain the responsibility of J. D. \"Pepper\" Petersen.": "This private experimental draft was prepared with AI assistance for restructuring, line editing, and production workflow. The thesis, judgment, architecture, and final editorial decisions remain those of J. D. \"Pepper\" Petersen.",
    "This book begins from a practical concern that has followed autonomous systems from the field into public institutions: action is moving faster than the structures that authorize it. A system that can sense, decide, coordinate, and act at machine speed cannot be governed only by after-the-fact policy, paper controls, or human review that arrives after consequence.": "This book begins with a problem I have seen in the field, in regulated markets, and in institutions trusted with real consequence: action is now moving faster than the structures that authorize it. Once a system can sense, decide, coordinate, and act at machine speed, policy that arrives after consequence is no longer governance. It is paperwork at the scene.",
    "The G-Plane is my attempt to describe a runtime architecture for that problem. The question is not whether autonomous systems can act, but how human authority remains present inside action itself: who may authorize it, where authority stops, what evidence survives, how escalation works, and how an institution can later explain what happened without pretending the machine was sovereign.": "The G-Plane is my answer to that problem. It does not ask whether machines can act. They can. It asks the harder question: how does human authority remain present inside the act itself? Who authorized it? Where does the authority stop? What evidence survives? Who may revoke it? What happens when the network goes dark? How does an institution explain the act afterward without pretending the machine became sovereign?",
    "Infrastructure is civilization made operational.": "Infrastructure is where civilization stops being an idea and becomes an operating condition.",
    "Governance must become runtime architecture.": "Governance has to move into the runtime.",
    "The central claim of this architecture is simple: no consequential action should execute unless it remains admissible under a valid authority chain at the exact moment execution occurs.": "The central claim is simple enough to be dangerous: no consequential machine action should execute unless it can show valid authority at the exact moment consequence becomes real.",
    "Artificial intelligence governance discussions frequently focus on model behavior, alignment, bias, explainability, or post-hoc auditing. Those discussions are necessary, but incomplete.": "Most AI governance talk begins in the wrong place. It talks about model behavior, alignment, bias, explainability, and post-hoc audit. Those subjects are necessary. They are not enough.",
    "At that moment governance must stop being descriptive. It must become enforceable.": "At that moment governance can no longer describe the world. It has to bind it.",
    "The solution cannot therefore be additional paperwork, larger monitoring systems, or more extensive compliance reporting. The solution must be architectural.": "The answer is not more paperwork, larger dashboards, or compliance reports with better adjectives. The answer is architecture.",
    "The modern world runs on invisible systems.": "The modern world runs on systems most people never see until they fail.",
    "Governance itself becomes infrastructure.": "At that point governance itself becomes infrastructure.",
    "For me the problem first became visible not through abstract artificial intelligence theory but through practical interaction with autonomous systems operating in physical environments.": "For me, the problem did not begin in an AI seminar. It began around physical systems, real airspace, weather, telemetry, liability, and machines that did not care whether the paperwork was elegant.",
    "Back in 2012, when Greg Heide and I were building Big Sky UAV in Montana, the commercial drone industry barely existed.": "Back in 2012, when Greg Heide and I were building Big Sky UAV in Montana, the commercial drone industry barely existed in the form people know now.",
    "The naive response to this problem is often to assume that stronger monitoring solves the issue.": "The lazy answer is monitoring.",
    "It does not, monitoring observes, and governance constrains.": "It does not. Monitoring observes. Governance constrains.",
    "Power does not become legitimate because it is intelligent. Power becomes governable only when it can show its authority before consequence occurs.": "Power does not become legitimate because it is intelligent. Power becomes governable only when it can show its authority before consequence occurs.",
}

STYLE_REWRITES = [
    ("The Governance Plane", "the G-Plane"),
    ("Governance Plane", "G-Plane"),
    ("autonomous infrastructure systems", "autonomous infrastructure"),
    ("infrastructure consequence", "consequence"),
    ("institutional legitimacy continuity", "legitimacy continuity"),
    ("runtime governance infrastructure", "runtime governance"),
    ("governance continuity", "governance continuity"),
    ("This distinction", "The distinction"),
    ("This separation", "The separation"),
    ("This capability", "The capability"),
    ("This requirement", "The requirement"),
    ("therefore becomes", "becomes"),
    ("therefore preserves", "preserves"),
    ("therefore treats", "treats"),
    ("therefore requires", "requires"),
    ("therefore separates", "separates"),
]

PARAGRAPH_REPLACEMENTS = {
    "The Governance Plane exists because autonomous systems operate inside environments where infrastructure consequence can occur faster than human supervisory intervention. Governance is not merely a record of what happened after the fact. It is the authority structure that must be present before irreversible consequence occurs.": "The G-Plane exists because machines can now reach consequence faster than human supervision can arrive. Governance is not the record we assemble afterward. It is the authority structure that must already be present when the act crosses into the world.",
    "Infrastructure systems may execute only when a valid authority chain remains active at execution time. That chain must remain continuous from constitutional origin to operational delegation to infrastructure consequence.": "Execution is legitimate only when the authority chain is alive at the moment of action. That chain must run from constitutional origin, through operational delegation, all the way to consequence.",
    "Governance artifacts must become executable runtime structures rather than static institutional documents. Systems operating at machine speed cannot depend on human interpretation at the moment of execution.": "Governance artifacts must become executable structures. A machine-speed system cannot wait for a human being to interpret a policy memo at the edge of execution.",
    "Distributed infrastructure often crosses operational and institutional boundaries. The architecture separates constitutional legitimacy, sovereign governance domains, infrastructure environments, delegated authority, and execution admissibility so that power does not dissolve into simple identity management.": "Distributed infrastructure crosses institutional boundaries as a matter of course. The architecture separates constitutional legitimacy, sovereign domain, local infrastructure, delegated authority, and admissibility so power does not collapse into mere login credentials.",
    "Standing permissions become dangerous in autonomous systems. Warrants provide single-use, execution-bound authority artifacts that are consumed at the execution boundary rather than persisting as broad ambient permission.": "Standing permission is dangerous when systems act continuously. Warrants turn authority into a single-use instrument, bound to the act and consumed at the boundary.",
    "Governed infrastructure must leave reconstructable proof of authority, constraint, admissibility, and consequence. The Governance Evidence Ledger preserves the institutional and technical conditions under which action occurred.": "Governed infrastructure must leave proof: authority, constraint, admissibility, and consequence. The Governance Evidence Ledger preserves the conditions under which power crossed the line.",
}


def has_image(paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def soften_sentence_rhythm(text: str) -> str:
    original = text
    for old, new in TARGETED_REWRITES.items():
        if text.strip() == old:
            return new
        text = text.replace(old, new)
    if original.strip() in PARAGRAPH_REPLACEMENTS:
        return PARAGRAPH_REPLACEMENTS[original.strip()]

    # Vary only body prose. Keep terms intact but reduce courthouse marble.
    if len(text) > 180:
        for old, new in STYLE_REWRITES:
            text = text.replace(old, new)

    text = text.replace("The architecture attempts to", "The architecture tries to")
    text = text.replace("The architecture exists to", "The architecture exists to")
    text = text.replace("The distinction is practical, not decorative.", "That is not decoration. It is the whole point.")
    text = text.replace("This is where the architecture becomes inspectable.", "This is where the architecture can be inspected.")
    text = text.replace("The general rule is simple:", "The rule is simple:")
    text = text.replace("That continuity is carried by a small set of primitives:", "The machinery is carried by a small set of primitives:")
    text = text.replace("admissible now", "admissible now")
    text = re.sub(r"\s+([,.;:?!])", r"\1", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    return text


def add_private_preface(doc: Document) -> None:
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Publication Notices":
            anchor = paragraph
            break
    if anchor is None:
        return
    preface_heading = anchor.insert_paragraph_before("Private Experimental Note", style="Heading 1")
    preface_heading.paragraph_format.page_break_before = True
    preface = anchor.insert_paragraph_before(
        "This version is not the publication copy. It is an authorial experiment: sharper in voice, more direct in argument, and closer to the operational experience that produced the G-Plane thesis. The architecture remains the same. The wager is style, not substance.",
        style="First Paragraph",
    )
    for run in preface.runs:
        run.italic = True


doc = Document(str(DOCX))
changed = 0
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == "Private Experimental Note":
        continue
    if has_image(paragraph):
        continue
    text = paragraph.text
    if not text.strip():
        continue
    new = soften_sentence_rhythm(text)
    if new != text:
        set_text(paragraph, new)
        changed += 1

add_private_preface(doc)

for paragraph in doc.paragraphs[:8]:
    if paragraph.text.strip() in {"POWER MUST SHOW ITS WARRANT", "The G-Plane Architecture for Governable Autonomy", 'J. D. "Pepper" Petersen', "Private Experimental Author's Draft | June 2026"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.text.strip() == "POWER MUST SHOW ITS WARRANT":
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(30)

doc.core_properties.title = "Power Must Show Its Warrant: The G-Plane Architecture"
doc.core_properties.author = 'J. D. "Pepper" Petersen'
doc.core_properties.comments = "Private experimental authorial rewrite. Not the publication copy and not linked from the site."
doc.save(str(DOCX))
print(f"changed={changed} output={DOCX}")
