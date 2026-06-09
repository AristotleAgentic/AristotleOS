from pathlib import Path
import re
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Scientific-Edition.docx"
TARGET = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"

shutil.copyfile(SOURCE, TARGET)
doc = Document(str(TARGET))


BROKEN_TEXT = {
    " untime ": " runtime ",
    " untime.": " runtime.",
    " untime,": " runtime,",
    " ully ": " fully ",
    " dvisory": " advisory",
    " nfrastructural": " infrastructural",
    " onsitutional": " constitutional",
    " nstitutional": " institutional",
    " nfrastructure": " infrastructure",
    " overnance": " governance",
    " uthority": " authority",
    " dmissibility": " admissibility",
    " hether": " whether",
    " hat happened": " what happened",
    " hat action": " what action",
    "Governance Invariantsdeterministic": "Governance Invariants: deterministic",
    "Meta Authority Envelopesconstitutional": "Meta Authority Envelopes: constitutional",
    "Wardsprotected": "Wards: protected",
    "Domainslocal": "Domains: local",
    "Envelopesdelegated": "Envelopes: delegated",
    "Registersmachine": "Registers: machine",
    "Warrantsaction": "Warrants: action",
    "Gatesexecution": "Gates: execution",
    "Gatershardware": "Gaters: hardware",
    "Ledgerscryptographic": "Ledgers: cryptographic",
}

PHRASE_REWRITES = [
    (r"\bThe Governance Plane therefore\b", "The Governance Plane"),
    (r"\bThe Governance Plane increasingly\b", "The Governance Plane"),
    (r"\bThe revised Governance Plane now assumes\b", "The revised architecture assumes"),
    (r"\bThe revised Governance Plane therefore\b", "The revised architecture"),
    (r"\bAutonomous infrastructure systems increasingly\b", "Autonomous infrastructure systems"),
    (r"\bInfrastructure systems increasingly\b", "Infrastructure systems"),
    (r"\bAutonomous systems increasingly\b", "Autonomous systems"),
    (r"\bsystems increasingly\b", "systems"),
    (r"\binfrastructure systems increasingly\b", "infrastructure systems"),
    (r"\bincreasingly increasingly\b", "increasingly"),
    (r"\btherefore therefore\b", "therefore"),
    (r"\bsubstantially strengthens\b", "strengthens"),
    (r"\bsignificantly strengthens\b", "strengthens"),
    (r"\bfundamentally strengthens\b", "strengthens"),
    (r"\bsignificantly shaped\b", "shaped"),
    (r"\bfundamentally shaped\b", "shaped"),
    (r"\bThis operational reality significantly shaped\b", "This operational reality shaped"),
    (r"\bThis layered continuity significantly strengthens\b", "This layered continuity strengthens"),
    (r"\bThe distinction may ultimately become one of\b", "The distinction may become one of"),
    (r"\bmay ultimately become\b", "may become"),
    (r"\bone of the defining operational requirements\b", "a defining operational requirement"),
]

FILLER_SENTENCES = [
    r"The distinction significantly matures the architecture\.",
    r"This distinction significantly matures the architecture\.",
    r"The distinction substantially matures the architecture\.",
    r"The transition significantly matures the architecture\.",
    r"This capability significantly matures the architecture\.",
    r"This capability substantially strengthens the architecture\.",
    r"This layered continuity substantially strengthens operational survivability\.",
    r"This layered continuity significantly strengthens deployability\.",
    r"This operational realism significantly strengthens the architecture\.",
]


def style_name(paragraph):
    return paragraph.style.name if paragraph.style else "Normal"


def is_heading(paragraph):
    return style_name(paragraph).startswith("Heading")


def has_image(paragraph):
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def normalize(text):
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"([.!?])\s+([a-z])", lambda m: f"{m.group(1)} {m.group(2).upper()}", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\.\s+\.", ".", text)
    text = text.replace(". But ", ". But ")
    return text.strip()


def polish_text(text):
    original = text
    for old, new in BROKEN_TEXT.items():
        text = text.replace(old, new)
    for pattern, replacement in PHRASE_REWRITES:
        text = re.sub(pattern, replacement, text)
    for pattern in FILLER_SENTENCES:
        text = re.sub(r"\s*" + pattern + r"\s*", " ", text)

    text = text.replace("systems. Rather than", "systems rather than")
    text = text.replace("rather than traditional governance hierarchy", "rather than a conventional governance hierarchy")
    text = text.replace("infrastructure. Rather than", "infrastructure rather than")
    text = text.replace("hierarchy. Rather than", "hierarchy rather than")
    text = text.replace("administration. Rather than", "administration rather than")
    text = text.replace("architecture. Rather than", "architecture rather than")
    text = text.replace("control. Rather than", "control rather than")
    text = text.replace("enforcement. Rather than", "enforcement rather than")
    text = text.replace("The Governance Plane separates:", "The Governance Plane separates")
    text = text.replace("The Governance Plane preserves:", "The Governance Plane preserves")
    text = text.replace("The Governance Plane governs:", "The Governance Plane governs")
    text = text.replace("That continuity requires:", "That continuity requires ")
    text = text.replace("not merely.", "not merely")
    text = text.replace("not artificial intelligence infrastructure, constitutional execution infrastructure", "not artificial intelligence infrastructure but constitutional execution infrastructure")
    return normalize(text) if text != original else original


def fix_front_matter():
    figure_anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Figure Plates":
            figure_anchor = paragraph
            break
    if figure_anchor is None:
        return

    figure_element = figure_anchor._element
    for paragraph in list(doc.paragraphs):
        if paragraph._element is figure_element:
            break
        delete_paragraph(paragraph)

    front_matter = [
        ("THE G-PLANE ARCHITECTURE", "Title"),
        ("Governance Infrastructure for Autonomous Systems", "Subtitle"),
        ("Including Wards, Warrants, Authority Routing, Evidence Ledgers, Governance Kernels, and Execution Control for Autonomous Infrastructure", "Normal"),
        ('J. D. "Pepper" Petersen', "Normal"),
        ("Aristotle Agentic Publication Edition | June 2026", "Normal"),
        ("Publication Notices", "Heading 1"),
        ('Copyright (c) 2026 J. D. "Pepper" Petersen. All rights reserved. This publication is issued by Aristotle Agentic as a research and governance-architecture work concerning autonomous systems, institutional authority, warrants, wards, evidence ledgers, governance kernels, and execution control.', "Normal"),
        ("The work is provided for research, education, policy, technical, and institutional discussion. It should not be treated as legal, financial, engineering-safety, procurement, export-control, aviation, insurance, or regulatory-compliance advice. Institutions applying these ideas should obtain appropriate professional review for their own legal duties, operating environments, and risk posture.", "Normal"),
        ("AI Assistance Disclosure", "Heading 1"),
        ('This publication package was prepared with the assistance of AI tools for formatting, readability editing, publication packaging, and production workflow. The underlying thesis, substantive judgment, research direction, and final publication decisions remain the responsibility of J. D. "Pepper" Petersen.', "Normal"),
        ("Rights and Citation", "Heading 1"),
        ('No part of this publication may be reproduced, stored in a retrieval system, or transmitted in any form or by any means without prior written permission, except for brief quotations, citation, scholarship, review, or other uses permitted by law. Suggested citation: Petersen, J. D. "Pepper". The G-Plane Architecture: Governance Infrastructure for Autonomous Systems. Aristotle Agentic, Publication Edition, 2026.', "Normal"),
        ("Author's Note on Scope", "Heading 1"),
        ("This book begins from a practical concern that has followed autonomous systems from the field into public institutions: action is moving faster than the structures that authorize it. A system that can sense, decide, coordinate, and act at machine speed cannot be governed only by after-the-fact policy, paper controls, or human review that arrives after consequence.", "First Paragraph"),
        ("The G-Plane is my attempt to describe a runtime architecture for that problem. The question is not whether autonomous systems can act, but how human authority remains present inside action itself: who may authorize it, where authority stops, what evidence survives, how escalation works, and how an institution can later explain what happened without pretending the machine was sovereign.", "Normal"),
    ]

    for text, style in front_matter:
        figure_anchor.insert_paragraph_before(text, style=style)


def merge_orphan_rather_than():
    merged = 0
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if paragraph._element is None or has_image(paragraph) or is_heading(paragraph):
            continue
        text = paragraph.text.strip()
        if not text.lower().startswith("rather than "):
            continue
        previous = None
        for prior in reversed(paragraphs[:idx]):
            if prior._element is not None and prior.text.strip() and not is_heading(prior) and not has_image(prior):
                previous = prior
                break
        if previous is None:
            continue
        prior_text = previous.text.strip().rstrip(".")
        combined = normalize(f"{prior_text} {text}")
        set_text(previous, combined)
        delete_paragraph(paragraph)
        merged += 1
    return merged


def insert_after(paragraph, text, style=None):
    new_paragraph = paragraph.insert_paragraph_before(text, style=style)
    paragraph._p.addprevious(new_paragraph._p)
    return new_paragraph


def fix_design_principles():
    heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Design Principles of the Governance Plane":
            heading = paragraph
            break
    if heading is None:
        return

    current = []
    capture = False
    for paragraph in list(doc.paragraphs):
        if paragraph._element is heading._element:
            capture = True
            continue
        if not capture:
            continue
        if is_heading(paragraph):
            break
        if paragraph.text.strip():
            current.append(paragraph)

    if not current:
        return

    principles = [
        ("Principle 1 — Governance must bind before consequence.", "The Governance Plane exists because autonomous systems operate inside environments where infrastructure consequence can occur faster than human supervisory intervention. Governance is not merely a record of what happened after the fact. It is the authority structure that must be present before irreversible consequence occurs."),
        ("Principle 2 — Authority must remain continuous.", "Infrastructure systems may execute only when a valid authority chain remains active at execution time. That chain must remain continuous from constitutional origin to operational delegation to infrastructure consequence."),
        ("Principle 3 — Governance must be machine enforceable.", "Governance artifacts must become executable runtime structures rather than static institutional documents. Systems operating at machine speed cannot depend on human interpretation at the moment of execution."),
        ("Principle 4 — Sovereignty must remain explicit.", "Distributed infrastructure often crosses operational and institutional boundaries. The Governance Plane separates constitutional legitimacy, sovereign governance domains, infrastructure environments, delegated authority, and execution admissibility so that power does not dissolve into simple identity management."),
        ("Principle 5 — Execution authority must be exhaustible.", "Standing permissions become dangerous in autonomous systems. Warrants provide single-use, execution-bound authority artifacts that are consumed at the execution boundary rather than persisting as broad ambient permission."),
        ("Principle 6 — Governance must produce evidence.", "Governed infrastructure must leave reconstructable proof of authority, constraint, admissibility, and consequence. The Governance Evidence Ledger preserves the institutional and technical conditions under which action occurred."),
    ]

    anchor = heading
    for title, body in reversed(principles):
        p = anchor.insert_paragraph_before(title, style="Heading 2")
        heading._p.addnext(p._p)
        anchor = p
        q = anchor.insert_paragraph_before(body, style="Body Text")
        anchor._p.addnext(q._p)
        anchor = q

    for paragraph in current:
        if paragraph._element is not None:
            delete_paragraph(paragraph)


def line_edit():
    changed = 0
    for paragraph in doc.paragraphs:
        if paragraph._element is None or is_heading(paragraph) or has_image(paragraph):
            continue
        text = paragraph.text
        if not text.strip():
            continue
        new = polish_text(text)
        if new != text:
            set_text(paragraph, new)
            changed += 1
    return changed


fix_front_matter()
changed = line_edit()
merged = merge_orphan_rather_than()
changed += line_edit()
fix_design_principles()

doc.core_properties.title = "The G-Plane Architecture: Governance Infrastructure for Autonomous Systems"
doc.core_properties.author = 'J. D. "Pepper" Petersen'
doc.core_properties.last_modified_by = "Aristotle Agentic"
doc.core_properties.comments = "Final candidate editorial pass preserving figures and architecture while improving prose flow, front matter, and damaged text."

doc.save(str(TARGET))
print(f"{TARGET} changed={changed} merged_orphans={merged} paragraphs={len(doc.paragraphs)}")
