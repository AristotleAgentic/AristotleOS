from pathlib import Path
import re
import sys

from docx import Document


target = Path(sys.argv[1])
doc = Document(str(target))

replacements = [
    (
        "Governance can no longer remain external to infrastructure operations. Governance must become part of infrastructure execution.",
        "Governance can no longer remain external to infrastructure operations; it has to become part of infrastructure execution.",
    ),
    (
        "Infrastructure environments rarely fail gracefully. They degrade, communications become intermittent, and telemetry becomes uncertain. Latency increases. Distributed systems drift out of synchronization. Local operators improvise. Operational priorities shift suddenly. Autonomous systems continue operating anyway because infrastructure itself cannot simply stop functioning.",
        "Infrastructure environments rarely fail gracefully: communications become intermittent, telemetry becomes uncertain, latency rises, distributed systems drift out of synchronization, local operators improvise, and operational priorities shift suddenly. Autonomous systems continue operating anyway because infrastructure itself cannot simply stop functioning.",
    ),
    (
        "Under traditional human systems constitutional violations are often discovered after execution. Courts review, investigations occur, and appeals emerge. Liability is assigned. But autonomous systems increasingly operate at machine velocity. Post hoc correction may become structurally insufficient. This realization is central. Once autonomous systems can release funds, deny access, coordinate infrastructure, allocate resources, direct physical systems, influence legal outcomes, mediate civic participation, and shape information environments.",
        "Under traditional human systems, constitutional violations are often discovered after execution: courts review, investigations proceed, appeals emerge, and liability is assigned. Autonomous systems change that timing. When systems can release funds, deny access, coordinate infrastructure, allocate resources, direct physical systems, influence legal outcomes, mediate civic participation, and shape information environments at machine velocity, post-hoc correction may become structurally insufficient.",
    ),
    (
        "This introduces an important principle: legitimacy is not static; it is continuously recomputed. This concept increasingly defines the Governance Plane. The architecture therefore increasingly resembles constitutional runtime. This is not metaphorical. The system literally computes whether authority remains constitutionally valid before execution may proceed. This becomes one of the architecture’s most important claims.",
        "This introduces an important principle: legitimacy is not static; it is continuously recomputed. The idea increasingly defines the Governance Plane and pushes the architecture toward constitutional runtime. That is not a metaphor. The system literally computes whether authority remains constitutionally valid before execution may proceed, which becomes one of the architecture’s central claims.",
    ),
    (
        "This distinction becomes profound. Civilization historically survived not because humans lacked power. Civilization survived because societies slowly constructed mechanisms that constrained power. Courts, constitutions, and due process. Separation of powers, civilian control, and checks and balances. Auditability, revocation, and institutional legitimacy.",
        "The distinction is profound. Civilization did not survive because humans lacked power; it survived because societies slowly constructed mechanisms that constrained power: courts, constitutions, due process, separation of powers, civilian control, checks and balances, auditability, revocation, and institutional legitimacy.",
    ),
    (
        "This becomes one of the architecture’s deepest warnings.",
        "That warning runs through the architecture.",
    ),
    (
        "This capability substantially strengthens the architecture.",
        "That capability strengthens the architecture.",
    ),
    (
        "This requirement strengthens the architecture significantly.",
        "The requirement strengthens the architecture.",
    ),
    (
        "This distinction changes everything.",
        "The difference changes the frame.",
    ),
    (
        "This shift changes the nature of governance itself.",
        "The shift changes the nature of governance itself.",
    ),
]


def polish(text):
    original = text
    for old, new in replacements:
        text = text.replace(old, new)

    text = text.replace("This realization materially sharpens", "The realization sharpens")
    text = text.replace("This realization significantly advances", "The realization advances")
    text = text.replace("This realization significantly strengthens", "The realization strengthens")
    text = text.replace("This capability may ultimately become", "That capability may become")
    text = text.replace("This concept increasingly defines", "The concept increasingly defines")
    text = text.replace("This distinction sits", "The distinction sits")
    text = text.replace("This distinction may become", "The distinction may become")
    text = text.replace("This separation", "The separation")
    text = text.replace("This transition", "The transition")
    text = text.replace("This architecture", "The architecture")
    text = text.replace("This realization", "The realization")
    text = text.replace("This distinction", "The distinction")

    # Remove a common stutter after replacement.
    text = re.sub(r"\bThe architecture therefore increasingly resembles\b", "The architecture increasingly resembles", text)
    text = re.sub(r"\btherefore therefore\b", "therefore", text, flags=re.I)
    return text if text != original else original


changed = 0
for paragraph in doc.paragraphs:
    text = paragraph.text
    new = polish(text)
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)
        changed += 1

doc.save(str(target))
print(f"changed={changed}")
