from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"

REPLACEMENTS = {
    " ay consequence": " may consequence",
    " rocedural": " procedural",
    " ommentary": " commentary",
    " ontrol": " control",
    " oes this actor": " does this actor",
    " oes ": " does ",
    " ho ": " who ",
    " hat ": " what ",
    " nder ": " under ",
    "Authority Envelope→": "Authority Envelope →",
    "Envelope→": "Envelope →",
    "Domain→": "Domain →",
    "Warrant→": "Warrant →",
    "Ward→": "Ward →",
    "Gate→": "Gate →",
    "authority generally? ": "authority generally? ",
    "access?.": "access?",
    "legitimate?.": "legitimate?",
    "implicated?.": "implicated?",
    "conditions?.": "conditions?",
    "environment?.": "environment?",
    "execution?.": "execution?",
    "domains?.": "domains?",
    "acting?.": "acting?",
    "delegated?.": "delegated?",
    "exist?.": "exist?",
    "now?.": "now?",
}


doc = Document(str(DOCX))
changed = 0
for paragraph in doc.paragraphs:
    text = paragraph.text
    new = text
    for old, replacement in REPLACEMENTS.items():
        new = new.replace(old, replacement)
    new = re.sub(r"\s+([,.;:?!])", r"\1", new)
    new = re.sub(r"\s{2,}", " ", new).strip()
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)
        changed += 1

doc.save(str(DOCX))
print(f"changed={changed}")
