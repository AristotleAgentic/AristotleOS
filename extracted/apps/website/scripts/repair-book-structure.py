from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"

READER_MAP_TEXTS = [
    "The Governance Plane is easiest to read as a chain from legitimacy to consequence. The book begins with the control problem, then builds the primitives that carry authority toward execution, then tests those primitives under failure, federation, emergency, and physical consequence.",
    "The primitives have distinct jobs. Meta Authority Envelopes establish who may create authority. Wards define the protected context on whose behalf authority exists. Authority Domains locate the infrastructure environment where consequence will occur. Authority Envelopes delegate bounded operational scope. Governance Invariants make constraints executable. Runtime Registers hold live state. Warrants carry action-specific authority. Commit Gates admit or refuse execution. Physical Invariant Gaters keep hard consequences unreachable. Evidence Ledgers preserve the chain after action.",
    "Several boundary terms recur, but they are not synonyms. Admissibility is the judgment that an act may proceed now. The admissibility state is the live condition set used to make that judgment. The Commit Gate is the enforcement mechanism. The sovereign commit boundary is the constitutional place that must be protected. The Last Boundary is the final image of the same crossing: possibility becoming consequence.",
    "Read the architecture in that order and the book becomes less a list of concepts than a single motion: authority begins above the machine, narrows as it approaches action, proves itself at the boundary, and leaves evidence behind.",
]


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def insert_before(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    new_para.style = style
    return new_para


def insert_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    new_para.style = style
    return new_para


def repair_reader_map(doc: Document) -> bool:
    changed = False
    reader_starts = tuple(text[:42] for text in READER_MAP_TEXTS)
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if paragraph._element is not None and (text in READER_MAP_TEXTS or text.startswith(reader_starts)):
            delete_paragraph(paragraph)
            changed = True
        elif paragraph._element is not None and text == "How to Read This Architecture":
            delete_paragraph(paragraph)
            changed = True

    target = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "Design Principles of the Governance Plane":
            target = paragraph
            break
    if target is None:
        return changed

    heading = insert_before(target, "How to Read This Architecture", "Heading 1")
    for text, style in reversed(
        [(READER_MAP_TEXTS[0], "First Paragraph")]
        + [(text, "Body Text") for text in READER_MAP_TEXTS[1:]]
    ):
        insert_after(heading, text, style)
    return True


def renumber_chapters(doc: Document) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 2" and re.fullmatch(r"Chapter \d+", text):
            count += 1
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(f"Chapter {count}")
    return count


def main() -> None:
    doc = Document(str(DOCX))
    reader_map = repair_reader_map(doc)
    chapters = renumber_chapters(doc)
    doc.core_properties.comments = "Repaired reader map placement and renumbered chapter markers after structural pass."
    doc.save(str(DOCX))
    print(f"reader_map_repaired={reader_map} chapters_numbered={chapters} output={DOCX}")


if __name__ == "__main__":
    main()
