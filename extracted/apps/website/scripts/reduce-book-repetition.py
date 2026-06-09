from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "book" / "The-G-Plane-Architecture-Final-Candidate.docx"


CHAPTER_REWRITES = {
    "Foreword": [
        "Every age eventually discovers that its tools have become institutions. The steam engine did not remain a machine; it reorganized labor, cities, transport, capital, and law. Electricity did not remain a utility; it changed the rhythm of public life. Computation did not remain an office instrument; it became the nervous system of commerce, government, logistics, medicine, and communication.",
        "Autonomous systems now introduce the next institutional problem. They do not merely advise human beings. They sense, infer, coordinate, and act inside environments where consequence may occur before ordinary review can arrive. The question is not whether these systems will be useful. They will be. The question is whether their usefulness will outrun the authority that makes action legitimate.",
        "The G-Plane is built from that concern. It treats governance not as sentiment, oversight, or post-hoc audit, but as operational architecture. Authority must be carried into execution. Constraint must bind before consequence. Evidence must survive after action. Human institutions must remain the source of legitimacy even when direct human approval is too slow for the operating environment.",
        "This book therefore asks one question in many forms: what must be true before an autonomous system may act? The answer is the architecture developed here: Wards, Warrants, Authority Domains, Authority Envelopes, Governance Invariants, Runtime Registers, Commit Gates, Physical Invariant Gaters, and Governance Evidence Ledgers. Together they express a single claim: power must show its warrant.",
    ],
    "The Control Problem": [
        "The control problem did not begin for me as an abstract AI problem. It appeared in physical systems, where the world does not pause while institutions decide what they meant. In 2012, while building Big Sky UAV in Montana, the commercial drone world was still raw. The systems were primitive compared with modern autonomous infrastructure, but the authority problem was already visible.",
        "A drone is not only software. It occupies airspace, crosses property, depends on telemetry, carries liability, and can move from safe operation to unsafe consequence in seconds. The question was never simply whether the aircraft could fly. The question was who had authorized the operation, what conditions made that authorization valid, what happened when those conditions changed, and who could later explain the action if something went wrong.",
        "Those are not ordinary software questions. They are runtime governance questions. Static permission is not enough when weather changes, GPS drifts, connectivity drops, mission scope shifts, or a new emergency authority enters the same operational space. Authority has to remain alive as the system moves.",
        "That realization became the first form of the G-Plane thesis. Autonomous systems require a governance layer that can evaluate authority at the moment of consequence. The control problem is therefore not solved by better dashboards or after-action records. It is solved by placing admissibility at the boundary where decision becomes action.",
    ],
    "Introducing the Governance Plane": [
        "The G-Plane is the runtime layer that asks whether a proposed consequential act may cross into the world. It does not replace the autonomous system. It does not try to make probabilistic intelligence deterministic. It places deterministic governance around the point where probabilistic output becomes infrastructure consequence.",
        "Figure 0.1 shows the full stack. The upper layers establish legitimacy and protected context. The middle layers translate authority into bounded operational delegation. The lower layers test execution, enforce physical limits, and preserve evidence. The point is not to slow every system to human speed. The point is to make authority machine-verifiable before action.",
        "This chapter therefore introduces the architecture as a stack, not as a philosophy. Meta Authority Envelopes define the root authority. Wards define the protected governance domain. Authority Domains bind authority to infrastructure locality. Authority Envelopes delegate operational scope. Governance Invariants compile constraints. Runtime Registers hold active state. Warrants carry action-specific authority. Commit Gates decide admissibility. Physical Invariant Gaters enforce hard limits. Governance Evidence Ledgers preserve reconstructable proof.",
        "The rest of the book develops these primitives in detail. The essential move happens here: governance is no longer outside the system looking back. It is inside the execution path looking forward.",
    ],
    "Governance Continuity Under Degraded Conditions": [
        "Governance becomes most important when operating conditions are least ideal. Networks partition, telemetry weakens, witnesses disagree, and revocation state may arrive late. Autonomous infrastructure cannot assume perfect synchronization, but it also cannot treat uncertainty as permission.",
        "The G-Plane handles degraded conditions through bounded continuity. Authority may narrow, execution windows may shrink, Commit Gates may require stronger evidence, and local domains may fail closed when legitimacy cannot be established. Degraded operation is therefore not an exception to governance. It is one of the conditions governance must be built to survive.",
        "The rule is simple: uncertainty may preserve limited safe continuity, but it may not create new authority. A system may continue only within pre-authorized bounds, under active constraints, with evidence sufficient for later reconstruction. Anything beyond that must pause, escalate, narrow, or refuse.",
    ],
    "Edge Governance and Local Admissibility": [
        "Edge governance is the local form of the same problem. A distant authority service may be unreachable, but the actuator, vehicle, network segment, clinic, facility, or field system still faces immediate conditions. Local admissibility determines what may happen at that edge without inventing new authority.",
        "The edge node therefore needs a small, current, machine-verifiable governance state: active Ward, Authority Domain, Authority Envelope, invariant set, Runtime Register snapshot, revocation vector, Warrant status, physical limits, and evidence obligation. The edge does not become sovereign because the network is down. It becomes responsible for enforcing the last valid bounded authority it can prove.",
        "Local admissibility is not degraded-governance theory repeated at smaller scale. It is the implementation surface where degraded governance becomes operational: what can this node prove, what can it safely do, what must it refuse, and what evidence must it preserve until reconciliation.",
    ],
    "Governance Meshes and Runtime Coordination": [
        "A Governance Mesh is the coordination fabric that lets governance state move through distributed systems without collapsing into a single central switch. It carries authority state, revocation state, warrant visibility, witness attestations, and evidence commitments across domains.",
        "The mesh is necessary because autonomous infrastructure rarely operates as one machine under one institution. It operates as fleets, agents, services, controllers, gateways, models, sensors, and physical systems distributed across administrative boundaries. Governance has to move with that topology.",
        "The mesh is therefore about reachability and propagation: which nodes know what, which attestations they trust, which revocations they have seen, which Warrants can still be honored, and which actions must wait for convergence.",
    ],
    "Runtime Synchronization and Governance State Convergence": [
        "Synchronization is the discipline that keeps distributed governance from becoming fiction. The question is not merely whether nodes exchange messages. The question is whether they converge on the governance state required to make action legitimate.",
        "Convergence has several dimensions: authority version, Ward state, revocation vector, invariant hash, telemetry freshness, Warrant status, witness confidence, and evidence continuity. A system may be operationally synchronized while governance is stale. That is not good enough for consequential action.",
        "Runtime synchronization therefore sets the conditions under which distributed nodes may continue, narrow, escalate, or refuse. It is the temporal logic of legitimacy in a system where no node sees everything at once.",
    ],
    "The Governance Evidence Ledger": [
        "The Governance Evidence Ledger is the memory of legitimacy. Ordinary logs describe behavior. The ledger reconstructs authority: who authorized the act, under which Ward, through which envelope, under what runtime state, with which Warrant, through which Commit Gate, and with what physical and evidentiary result.",
        "This distinction matters because autonomous consequence often fragments responsibility. A model recommended, an agent planned, a service routed, a controller executed, and an institution deployed. Without a ledger, the chain becomes blame without lineage. With a ledger, the act has an authority history.",
        "The ledger is therefore not decorative audit storage. It is part of the governance system itself. If authority cannot be reconstructed, it was not operationally governed in the first place.",
    ],
    "The Governance Evidence Chain and Forensic Continuity": [
        "Forensic continuity begins where the ledger becomes useful to an outside reviewer. It asks whether a regulator, insurer, court, operator, or affected institution can reconstruct the chain of legitimacy without trusting the autonomous system's own story.",
        "The evidence chain connects records across time: policy source, authority artifact, invariant compilation, Runtime Register state, Warrant issuance, Commit Gate decision, physical gate status, execution result, witness attestation, and reconciliation. Each link should be independently checkable enough to survive dispute.",
        "This chapter is therefore not a second ledger chapter. It is about evidentiary durability: how the record remains coherent after synchronization delay, system failure, contested action, incident review, insurance analysis, or legal challenge.",
    ],
    "Governance as Civilizational Infrastructure": [
        "Civilization depends on constraints that survive the ambition of its tools. Roads, grids, courts, markets, communications networks, and aviation systems became trustworthy only when power was bounded by institutions, procedures, evidence, and enforceable limits.",
        "Autonomous systems force the same transition at machine speed. If these systems are allowed to act merely because they are useful, practical authority will migrate into runtime infrastructure faster than institutions can name the transfer. The result will not look like a coup. It will look like convenience.",
        "The civilizational claim of the G-Plane is that legitimate authority must remain structurally present inside machine action. Governance becomes infrastructure because without it, infrastructure becomes government by execution.",
    ],
    "Governance Native Infrastructure": [
        "Governance-native infrastructure is the practical design response to that civilizational claim. It means systems are built with authority, admissibility, revocation, physical containment, and evidence as first-order runtime surfaces rather than afterthoughts.",
        "A governance-native system does not bolt compliance onto an autonomous stack after deployment. It carries Wards, Authority Domains, Authority Envelopes, invariants, Runtime Registers, Warrants, Commit Gates, and ledgers as operational dependencies.",
        "The design standard is straightforward: if a consequential action cannot show its authority, the system should not be able to execute it. Governance-native infrastructure makes that refusal a property of the system rather than a hope placed on its operators.",
    ],
    "Runtime Authority Routing and Governance Path Selection": [
        "Authority routing determines which governance path applies before action. In distributed infrastructure, the relevant authority may depend on Ward, location, domain, mission, emergency state, revocation posture, physical consequence, and federation rules.",
        "This is different from revocation. Revocation decides whether authority has narrowed or disappeared. Routing decides which authority chain must be consulted in the first place. A system that routes authority incorrectly may execute under the wrong sovereign context even if every later check behaves correctly.",
        "Governance path selection therefore becomes a runtime function. The system must locate the controlling Ward, identify the active Authority Domain, select the applicable envelope, evaluate relevant invariants, and determine whether a Warrant can be issued or honored.",
    ],
    "Constitutional Runtime Governance": [
        "Constitutional runtime governance is the claim that legitimacy has to be evaluated inside execution, not merely declared above it. The word constitutional does not mean that every deployment is a state. It means that every consequential system needs a root of authority that operational systems cannot create for themselves.",
        "The runtime side of the phrase is equally important. A constitution that cannot reach the execution boundary is only background language. Runtime governance asks whether root authority, protected domain, delegated scope, active conditions, revocation state, and evidence obligation remain valid at the moment of action.",
        "This chapter establishes the vertical logic of the architecture: authority begins above the machine, travels through defined governance artifacts, narrows as it approaches execution, and either becomes admissible at the Commit Gate or fails before consequence.",
    ],
    "Meta Authority Envelopes": [
        "A Meta Authority Envelope is the root instrument of the G-Plane. It defines who may create protected domains, issue authority, amend governance structure, revoke downstream authority, and approve the rules by which governance becomes executable.",
        "The MAE exists to prevent operational systems from becoming self-authorizing. Without a root authority layer, a sufficiently capable system can gradually treat configuration, access, optimization, or emergency exception as practical sovereignty. The MAE keeps ultimate authority outside the machine.",
        "In implementation, the MAE may resemble a signed charter, constitutional authority document, institutional trust anchor, governance root, or regulated authorization source. Its form can vary. Its function cannot: it establishes the legitimacy from which all downstream authority derives.",
    ],
    "From Concept to Mechanism": [
        "The concept becomes useful only when it becomes mechanism. The G-Plane does not rely on one large governance object. It separates authority into artifacts with different jobs so that each can be evaluated, revoked, inspected, and implemented without collapsing the whole system into vague policy.",
        "The mechanism chain is deliberately staged. Meta Authority Envelopes establish root legitimacy. Wards define protected context. Authority Domains bind governance to infrastructure locality. Authority Envelopes delegate operational scope. Governance Invariants compile constraints. Runtime Registers carry live state. Warrants authorize specific acts. Commit Gates decide admissibility. Physical Invariant Gaters enforce hard limits. Evidence Ledgers preserve proof.",
        "This chapter is the bridge from thesis to engineering. It explains why the architecture is decomposed into primitives before later chapters examine each primitive in detail.",
    ],
    "Authority Envelopes": [
        "An Authority Envelope is delegated power with boundaries. It tells a system what class of action may be considered, under what conditions, for how long, within which domain, and subject to which constraints. It is broader than a Warrant and narrower than root authority.",
        "The envelope is not the final permission to act. That distinction matters. An envelope defines the possibility space; a Warrant authorizes a specific proposed consequence inside that space. Without that separation, broad delegation turns into standing authority.",
        "Authority Envelopes therefore make delegation machine-readable while preserving the need for execution-bound admissibility. They are how institutions delegate operational capacity without surrendering the boundary where action becomes consequence.",
    ],
    "Warrants": [
        "A Warrant is the action-level proof that a consequential act may proceed now. It is not identity, not role, not general permission, and not a reusable token. It is execution-bound authority tied to a proposed act, a Ward, an Authority Domain, an envelope, active state, and a short execution window.",
        "The reason for the Warrant is simple: standing permissions are too loose for autonomous systems. A system may be generally authorized and still inadmissible in the moment. Telemetry may change, revocation may propagate, a Ward may narrow, or the environment may cross a physical threshold.",
        "The Warrant turns delegated authority into present-tense legitimacy. It must be checked at the Commit Gate, consumed or refused, and preserved in the evidence record. This is where authority stops being ambient and becomes accountable.",
    ],
    "Where the Problem First Became Visible": [
        "The Governance Plane did not begin as a clean theory. It began in operating environments where the world kept moving while authority was still being interpreted. In the early Big Sky UAV years, autonomous aircraft were modest by today's standards, but they exposed the same problem that now appears across agentic infrastructure: a system may be technically capable of acting before the institution has finished determining whether it should.",
        "Drone operations made the issue concrete. A mission could begin under clear authorization and become ambiguous in minutes because weather shifted, telemetry weakened, GPS confidence changed, a new public-safety priority appeared, or the aircraft crossed into a different operational context. None of those changes were merely engineering details. They altered the legitimacy of action.",
        "That experience shaped the architecture more than any abstract debate about AI. The early lesson was that authority has to remain alive while a system moves. It cannot sit behind the operation as a static credential. It must travel with the system, narrow when conditions narrow, and be capable of refusing action when the world has changed.",
        "The Ward concept eventually came from this same pressure. A wildfire response flight, for example, is not governed only by airspace. It is governed by emergency authority, landowner interests, communications constraints, public-safety obligations, insurance exposure, and institutional accountability. The Ward preserves the protected context on whose behalf the operation occurs, while Authority Domains describe the local infrastructure environments through which the system moves.",
        "This is why the G-Plane is a constitutional runtime architecture rather than a compliance wrapper. It was born from the practical fact that autonomous systems create consequence in live environments. Governance has to stand at that boundary, not arrive afterward with a report.",
    ],
    "Authority Domains": [
        "Authority Domains answer the locality question. Once a system is authorized in principle, the next question is where that authority is being exercised. Consequence never occurs in an abstract software space. It occurs in airspace, a clinic, a substation, a warehouse, a network segment, a public agency workflow, a field operation, or a communications corridor.",
        "Each domain carries its own conditions. Telemetry freshness, physical risk, jurisdiction, operational ownership, emergency posture, safety limits, synchronization state, and revocation visibility may differ from one domain to the next. A proposed act that is admissible inside one domain may become inadmissible seconds later in another.",
        "This is the distinction between Wards and Authority Domains. A Ward answers: on whose protected behalf does authority exist? An Authority Domain answers: inside which operational environment will consequence occur? The first preserves sovereign legitimacy. The second preserves local consequence context.",
        "The distinction is not academic. A healthcare system may contain pharmacy logistics, patient monitoring, facility automation, emergency coordination, and robotics domains under a broader patient-care Ward. An emergency response operation may contain airspace, radio, logistics, utility, and command domains under one or more emergency Wards. The architecture needs both ideas because legitimacy and locality do different work.",
        "Authority Domains also make degraded governance possible. If a central service is unreachable, the local domain can still enforce the last valid bounded authority it can prove. It may continue narrowly, require stronger evidence, or fail closed. What it may not do is invent new sovereignty because the network is down.",
    ],
    "Constitutional AI Versus Administrative AI": [
        "Administrative AI helps institutions process work. It drafts, summarizes, routes, classifies, searches, schedules, and recommends. These uses can be valuable, but they usually remain downstream of ordinary institutional authority. Constitutional AI, as used in this book, concerns a different threshold: the point at which machine systems participate in the formation, delegation, narrowing, or execution of authority itself.",
        "The difference is not intelligence. It is consequence. A document assistant may be powerful without becoming constitutionally significant. A system that can issue derived authority, alter escalation paths, create subordinate agents, or make consequential actions reachable has entered a more serious class of governance.",
        "The G-Plane is built for that second class. It asks whether authority has a valid root, whether sovereign context is preserved, whether delegation is bounded, whether admissibility can be tested at runtime, and whether evidence can reconstruct the act afterward. Administrative oversight is not enough once a system begins shaping the conditions under which power can be exercised.",
        "This boundary will become increasingly important as organizations adopt agentic systems. The institution may believe it is merely automating workflow while, in practice, it is allowing authority to migrate into infrastructure. The constitutional frame gives leaders a way to see that migration before it becomes normal.",
    ],
    "The Constitutional Boundary": [
        "The constitutional boundary is the line between ordinary operation and legitimate consequence. On one side, a system may compute, recommend, simulate, draft, or plan. On the other side, it changes the world under institutional authority. The Governance Plane exists because this boundary is where machine capability must be tested against legitimacy.",
        "A policy check asks whether a proposed act appears to violate a rule. Constitutional admissibility asks a harder question: does this system possess legitimate authority to create this category of consequence, in this context, at this moment, under this chain of delegation? That question cannot be answered by a credential alone.",
        "The boundary therefore requires more than permission. It requires a valid Ward, an applicable Authority Domain, a bounded Authority Envelope, active invariants, current registers, an unexpired Warrant, a Commit Gate decision, and evidence sufficient for later review. These are not decorative controls. They are the machinery that keeps execution subordinate to authority.",
        "When the boundary is weak, institutions may keep the language of control while losing control in practice. When it is strong, machine action can remain useful without becoming self-authorizing. The constitutional boundary is where that difference becomes operational.",
    ],
    "The Constitutional Layer": [
        "The constitutional layer governs the authority that governs action. Ordinary policy decides what a system may do. Constitutional governance decides who may define those permissions, amend them, delegate them, revoke them, or create new authority beneath them. That distinction becomes central once autonomous systems begin creating agents, composing workflows, and modifying operational structures.",
        "Without a constitutional layer, authority can drift invisibly. Configuration changes, model updates, emergency exceptions, automation shortcuts, and optimization pressure can gradually alter who effectively holds power. Nothing dramatic has to happen. The institution may simply discover that the real boundary of authority has moved.",
        "The Meta Authority Envelope exists to prevent that drift. It keeps amendment authority separate from operational authority. A logistics agent may reroute vehicles, but it may not grant itself jurisdiction, remove escalation rights, weaken audit guarantees, or bypass physical gates. A public-sector system may assist drafting, but it may not become the institution that authorizes the final act.",
        "This is not a preference for bureaucracy over speed. It is the old constitutional lesson translated into machine terms: the power to act and the power to redefine authority cannot safely collapse into the same mechanism. The constitutional layer gives that separation an executable form.",
    ],
    "Recursive Governance": [
        "Recursive governance begins when systems help govern other systems. Delegation itself is not new; governments, courts, militaries, firms, and agencies have always delegated. The new difficulty is speed, scale, and opacity. Machine systems can generate subordinate agents, compose execution paths, allocate authority-like permissions, and adapt governance routing faster than institutions can follow unless the chain is designed to remain visible.",
        "The danger is not recursion by itself. The danger is uncontrolled recursion. A local permission becomes a broader workflow. A workflow creates subordinate agents. Those agents federate across domains. Exceptions become routines. Escalation narrows. Over time, authority may expand without any single decision that looks like a constitutional event.",
        "The G-Plane responds by separating operational adaptation from constitutional amendment. Authority Envelopes define what a system may do. Meta Authority Envelopes define what a system may authorize. Wards keep recursive action attached to protected legitimacy. Warrants prevent delegated power from becoming permanent ambient authority. Evidence Ledgers preserve the lineage when the chain becomes too complex for memory alone.",
        "The result is a deliberate asymmetry. Operations may adapt quickly; legitimacy must mutate carefully. Execution may scale; authority propagation must remain bounded. Recursive systems are inevitable in agentic infrastructure. Recursive sovereignty is a design failure.",
    ],
    "The Constitutional Execution Chain": [
        "Civilization has always depended on chains of legitimacy. A judge acts under law. A commander acts under commission. A public agency acts under statute. A financial institution settles under recognized authority. Autonomous systems do not remove this requirement. They make it harder to preserve because execution can be delegated, composed, routed, and completed before a human institution can reconstruct the path.",
        "The constitutional execution chain is the continuous lineage connecting a consequential machine act to valid authority. It begins above the machine with a Meta Authority Envelope, passes through a Ward, narrows through Authority Domains and Authority Envelopes, becomes action-specific through a Warrant, and reaches the Commit Gate before consequence.",
        "This chain must be reconstructable. A system may succeed operationally and still fail constitutionally if no one can show whose authority governed the act, which protected context applied, which delegation permitted it, which conditions were active, and why execution was admitted. Capability is not enough; lineage is part of legitimacy.",
        "The chain also explains why evidence is not an afterthought. The Governance Evidence Ledger preserves the authority history of action. It allows regulators, courts, insurers, operators, and affected institutions to ask not only what happened, but whether the action remained legitimately reachable when it happened.",
    ],
    "Constitutional Failure Modes": [
        "A constitutional failure mode is not merely a software bug. It is a condition in which machine execution remains technically possible while legitimate authority has broken, narrowed, become ambiguous, or lost its evidence chain. These failures are dangerous precisely because the system may appear to be working.",
        "One failure mode is lineage break: the action proceeds, but the system can no longer prove its path back to valid authority. Another is authority inflation: delegated scope expands through workflow convenience until it exceeds the original grant. A third is revocation blindness: the system continues because it has not received, recognized, or honored narrowing authority.",
        "Other failures are local. A domain may operate on stale telemetry. A Warrant may be treated as reusable. A Commit Gate may rely on incomplete registers. A physical gate may fail open. A witness may attest to state it did not actually observe. In each case, the architecture has to distinguish ordinary operational failure from illegitimate consequence.",
        "The point of naming these modes is practical. Builders can test them. Institutions can require controls against them. Regulators can ask for evidence of refusal behavior. Insurance markets can price systems that can prove the difference between governed failure and inadmissible execution.",
    ],
    "Governance Compilation and Deterministic Runtime Enforcement": [
        "Human governance usually begins in language. Statutes, charters, contracts, orders, policies, safety standards, and institutional doctrine are written for interpretation. Autonomous infrastructure cannot interpret all of that language at the moment an actuator, transaction, route, or field system is about to create consequence. Something must translate the relevant authority into conditions a machine can test.",
        "Governance compilation performs that translation. It does not replace law, policy, or judgment. It extracts the enforceable conditions required for runtime control: who may act, under which Ward, inside which Authority Domain, within what scope, during what window, against what telemetry, under what physical limits, and with what evidence obligation.",
        "The compiled result is a Governance Invariant. An invariant is not the source of legitimacy. It is the operational expression of a legitimate rule at the boundary of action. This keeps the architecture from pretending that machines understand institutions merely because they can process text.",
        "Deterministic enforcement begins after compilation. Probabilistic systems may reason, plan, recommend, or optimize, but deterministic gates decide whether consequence is admissible. The distinction lets autonomy remain flexible while keeping irreversible action inside enforceable boundaries.",
    ],
    "Execution Timing, Commit Sequencing, and Consequence Windows": [
        "Governance becomes temporal as it approaches consequence. A system may be authorized at 10:00:00 and inadmissible at 10:00:03 because telemetry changed, a revocation arrived, a domain shifted, a physical threshold was crossed, or an emergency priority displaced the original mission.",
        "The G-Plane therefore treats execution as a sequence rather than a single instant. Authority is checked, state is refreshed, invariants are evaluated, a Warrant is validated, the Commit Gate decides, physical constraints are confirmed, and the evidence record is opened before the act crosses into infrastructure.",
        "A Consequence Window is the short interval during which that chain remains valid. It prevents authority from becoming stale while a system waits, retries, routes, or coordinates across distributed infrastructure. The window may shrink under degraded telemetry or expand only when the governing authority permits it.",
        "This chapter's contribution is narrow: timing is itself a governance surface. If legitimacy cannot survive the sequence from proposal to commit, the action should not proceed.",
    ],
    "The Governance Kernel": [
        "The Governance Kernel is the minimum runtime machinery that makes the architecture real. Policies, ledgers, charters, and diagrams may define governance, but the kernel is where a proposed consequential act is admitted or refused.",
        "The kernel sits at the convergence of active Ward, Authority Domain, Authority Envelope, invariant set, Runtime Register state, Warrant status, revocation visibility, telemetry freshness, timing window, and physical containment. It does not need to understand every institutional reason behind the rule. It needs to enforce the conditions that must hold before action.",
        "This is why the kernel resembles a commit engine more than a dashboard. It receives a proposed act and determines whether the act remains reachable under current authority. If the answer is no, execution stops before the system changes the world.",
        "The kernel also separates probabilistic intelligence from deterministic consequence. Models may remain adaptive. Planners may remain exploratory. Agents may remain useful. But the final transition into infrastructure must pass through a substrate that can say no.",
    ],
    "Revocation, Narrowing, and Dynamic Authority Control": [
        "Revocation is the proof that authority remains alive. If an institution cannot narrow, suspend, or withdraw delegated power before consequence propagates, then delegation has begun to resemble surrender.",
        "Autonomous infrastructure makes revocation difficult because authority state travels through distributed systems. Some nodes receive updates immediately. Others operate under delay, degraded connectivity, or partial visibility. The architecture cannot assume perfect synchronization, but it also cannot allow stale authority to become a license for new consequence.",
        "The G-Plane handles this through narrowing. Authority Envelopes may contract, Warrants may expire, Commit Gates may demand stronger evidence, physical gates may harden, and local domains may fail closed when revocation state is uncertain. The system does not move from fully authorized to fully dead in one brittle step. It moves through bounded states of reduced authority.",
        "Dynamic authority control is therefore the operational discipline of keeping delegated power subordinate as conditions change. It is not account administration. It is runtime legitimacy management.",
    ],
    "Governance Failure Modes and Graceful Degradation": [
        "Infrastructure rarely fails in a clean line. Sensors drift, links partition, witnesses disagree, clocks skew, operators improvise, and emergency priorities appear while systems are still running. A governance architecture that only works under perfect conditions is not an infrastructure architecture.",
        "Graceful degradation means the system loses authority carefully. As confidence declines, action windows shorten, admissibility thresholds rise, physical limits tighten, local autonomy narrows, and escalation becomes more likely. The design aim is not to keep every capability alive. It is to keep legitimacy alive.",
        "This gives operators a vocabulary beyond up or down. A domain may continue limited safe action under a valid cached Warrant. It may refuse high-consequence action until synchronization returns. It may preserve evidence for later reconciliation. It may enter emergency mode only through a defined authority path.",
        "The relevant test is whether the system fails toward bounded consequence. If uncertainty creates broader authority, the architecture has failed. If uncertainty narrows authority while preserving necessary safety and evidence, governance has survived degradation.",
    ],
    "Governance State Machines and Authority Lifecycles": [
        "Authority has a lifecycle. It is issued, activated, narrowed, suspended, delegated, synchronized, revoked, expired, reconciled, or recovered. Treating authority as a static permission hides the transitions where many failures begin.",
        "Governance State Machines make those transitions explicit. A Warrant can move from issued to active to consumed, expired, refused, or revoked. An Authority Envelope can activate, narrow, suspend, supersede, or terminate. A Ward can federate, isolate, enter emergency posture, or recover. Each transition carries evidence and conditions.",
        "This lifecycle view gives the architecture a way to handle real infrastructure behavior. Systems operate through latency, partial updates, changing telemetry, and shifting domains. The state machine tells the Commit Gate what authority exists now, not merely what was granted earlier.",
        "Forensic value follows from the same structure. When a later reviewer asks why a system acted, the answer is not only a log entry. It is the sequence of authority states that made the act admissible or should have caused refusal.",
    ],
    "Sovereignty, Jurisdiction, and the Fragmentation of Machine Authority": [
        "Autonomous systems will not enter a politically unified world. They will move through overlapping jurisdictions, private platforms, public agencies, allied command structures, local infrastructure, tribal and regional authority, regulated industries, and cross-border supply chains. Technical reach will often be transnational while legitimacy remains jurisdictional.",
        "That mismatch creates a governance problem. A system may possess operational capability across many domains while holding legitimate authority in only one. If it crosses the boundary without recognizing the difference, it creates legitimacy collision: two or more authority systems may disagree about who may act, who may revoke, whose evidence is sufficient, and which consequences are admissible.",
        "The G-Plane treats sovereignty as an operating condition rather than a nuisance. Wards preserve protected legitimacy. Authority Domains preserve local infrastructure context. Interdomain routing determines which authority path applies. Revocation and evidence must survive federation without merging sovereign control.",
        "This is bounded interoperability. Systems may coordinate across domains, but they may not dissolve the authority of the domains they enter. The future will not be governed by one global machine authority. It will require machinery that lets autonomous systems cooperate while keeping legitimacy local, visible, and interruptible.",
    ],
    "Wards and Sovereign Governance Domains": [
        "The Ward is the architecture's answer to a question ordinary permission systems usually avoid: on whose protected behalf does authority exist? That question becomes unavoidable once autonomous systems leave the world of software tasks and begin acting inside hospitals, airspace, energy systems, public agencies, emergency zones, financial networks, or other places where consequence belongs to people and institutions.",
        "A Ward is not a tenant, namespace, role, account, partition, or access-control group. Those structures organize systems. A Ward protects legitimacy. It names the sovereign or institutional context that gives delegated action its rightful source: patient care, emergency response, fiduciary obligation, public authority, critical infrastructure continuity, defense command, or another protected domain of consequence.",
        "This separates the Ward from the Authority Domain. The Ward tells the system whose protected interest is being served. The Authority Domain tells the system where the consequence will occur. A hospital may operate pharmacy logistics, patient monitoring, robotics, and facility automation domains inside a broader Patient Care Ward. A wildfire operation may contain aircraft, communications, utilities, logistics, and command domains inside an Emergency Response Ward.",
        "The distinction prevents operational cooperation from becoming accidental sovereignty. Systems may federate, exchange evidence, coordinate movement, or share telemetry without merging the authority under which they act. Every Authority Envelope, Warrant, Commit Gate decision, and evidence record must remain attached to a Ward so that later review can determine not only what happened, but under whose legitimate authority it happened.",
        "The Ward is therefore the constitutional perimeter of delegated machine action. It keeps authority from dissolving into topology, convenience, or technical reach. Without it, autonomous systems may appear authorized because they can operate. With it, they must remain authorized by the protected context that makes operation legitimate.",
    ],
    "Federated Governance and Interdomain Authority Routing": [
        "Federated governance begins where one institution's authority is no longer enough. Disaster response, telecommunications, autonomous logistics, distributed energy, aviation, healthcare, defense support, and cross-border infrastructure all require systems to coordinate across domains that do not share a single sovereign command.",
        "The governance problem is not merely coordination. Machines already coordinate well. The problem is legitimacy across coordination: which Ward governs this action, which host domain constrains it, whose revocation must be honored, which authority travels with the system, and which evidence must survive when several institutions participate in one consequence.",
        "Interdomain Authority Routing is the mechanism for answering those questions at runtime. Like network routing, it determines the valid path before something moves. Unlike network routing, the packet is not just information. It is delegated authority approaching consequence. The route must therefore preserve origin, host constraints, revocation state, admissibility limits, and evidence obligations.",
        "Federation must narrow authority rather than inflate it. A system should not gain power because it crosses more domains. It should usually become more constrained, because more sovereign interests are implicated. The admissible path is the intersection of legitimate authorities, not the most permissive union available.",
        "This is the practical role of Governance Meshes. They allow authority state, revocation, witness attestations, Warrant visibility, and evidence commitments to move across institutions without forcing those institutions into one hierarchy. The result is interoperability without surrender: systems cooperate, but sovereignty remains legible.",
    ],
    "The Constitutional Runtime and the Future of Machine Civilization": [
        "A civilization becomes durable when it learns how to bind power without stopping useful action. Law, procedure, chain of command, licensing, inspection, audit, and constitutional limits all emerged because raw capability was never enough. The autonomous era presses that old lesson into a new operating environment.",
        "Agentic systems increasingly sense, plan, negotiate, coordinate, and execute in places where human review cannot remain upstream of every act. If governance stays outside those pathways, institutions will retain the language of control while losing the boundary where control matters. The constitutional runtime is the attempt to move legitimate authority into that boundary.",
        "The term constitutional is not ornamental. It means the system must preserve a root of legitimacy, a protected context, delegated authority, present-tense constraints, revocation, and evidence before action reaches consequence. The term runtime is equally important: these conditions must be evaluated while the system is operating, not reconstructed only after damage has occurred.",
        "The future of machine civilization will not be decided only by model capability. It will be decided by whether powerful systems can remain subordinate to human institutions at the speed of execution. The Governance Plane offers one answer: make legitimacy operational, make authority testable, and make consequence refuseable.",
    ],
    "The Constitutional Machine": [
        "The constitutional machine is not a machine that replaces constitutional government. It is a machine whose power is shaped by constitutional constraints before it acts. That distinction is essential. The G-Plane does not ask software to become sovereign. It asks software to remain visibly subordinate to authority it did not create.",
        "Such a machine carries its limits with it. It knows the Ward under which it acts, the Authority Domain it has entered, the scope of its envelope, the state of revocation, the current register values, the Warrant it must present, and the evidence it must leave behind. These are not decorative labels. They are the conditions under which the machine may touch the world.",
        "This changes the meaning of autonomy. Autonomy no longer means unbounded self-direction. It means delegated discretion inside a lawful possibility space. The machine may optimize, plan, and coordinate, but it may not expand its own rightful power merely because expansion is efficient.",
        "A constitutional machine is therefore powerful and interruptible at the same time. It can act at machine speed while remaining answerable to institutions that move more slowly. That is the design problem this book has been circling from the beginning.",
    ],
    "The Age of Admissibility": [
        "The autonomous century will need a word stronger than compliance. Compliance looks backward and asks whether behavior matched a rule. Admissibility looks forward and asks whether this action may cross into consequence now.",
        "That change in tense is the change in governance. A system may have been compliant yesterday and inadmissible today. It may hold a credential and still lack authority. It may produce a correct prediction and still be barred from execution. It may be useful and still be outside its Ward.",
        "Admissibility brings authority, state, timing, telemetry, revocation, physical containment, and evidence into one decision. It is not a moral slogan. It is an operational judgment at the boundary where decision becomes action.",
        "This is why the architecture centers the Commit Gate. The Gate is where admissibility becomes real. Everything upstream prepares the question; everything downstream records the answer. In the age of admissibility, the decisive question is not can the system act, but may this consequence occur under legitimate authority at this moment?",
    ],
    "Constitutional Civilization": [
        "A constitutional civilization is not one that eliminates power. It is one that makes power answerable before it becomes irreversible. That has always been the civilizational bargain. Courts may coerce, armies may use force, agencies may regulate, banks may move value, and infrastructure operators may control essential systems, but each form of power must be tied to authority beyond itself.",
        "Autonomous systems threaten that bargain when practical authority migrates into runtime infrastructure without public recognition. The system routes, filters, allocates, approves, denies, escalates, or actuates. No constitution is repealed. No office is abolished. Yet the place where consequence becomes reachable moves away from the institution and into the machine.",
        "The G-Plane is an architecture against that drift. It does not stop autonomy. It gives autonomy a constitutional shape: Wards for protected context, Warrants for action-level authority, Gates for admissibility, physical constraints for hard limits, and ledgers for institutional memory.",
        "The claim is modest and severe at the same time. Machine systems may become part of civilization's operating fabric, but they may not become the source of their own legitimacy. Constitutional civilization survives autonomy only if authority remains visible where action becomes real.",
    ],
    "Governance as Infrastructure": [
        "Governance becomes infrastructure when it is no longer merely spoken about, written down, or audited afterward. It becomes infrastructure when systems depend on it in order to act.",
        "That is the design shift proposed here. A Warrant is not a policy memo. A Commit Gate is not a compliance meeting. A Runtime Register is not a report. A Physical Invariant Gater is not an aspiration. Each is part of the machinery through which authority reaches consequence.",
        "This shift will feel unfamiliar because many institutions still treat governance as overhead. In autonomous infrastructure, governance is closer to braking, routing, authentication, settlement, or fault isolation. It is a function the system must perform continuously to remain deployable.",
        "The practical implication is direct: future infrastructure buyers, regulators, insurers, and operators should ask not only whether a system is intelligent, but whether it is governable. Can it show authority? Can it narrow? Can it refuse? Can it preserve evidence? If not, governance has not become infrastructure; it remains a promise outside the machine.",
    ],
    "The Governed Future": [
        "The governed future will be built through ordinary institutional work. Standards will have to be drafted, audits designed, insurance models tested, procurement language changed, domain pilots run, and regulators taught to ask better questions. The architecture will not become real by being admired.",
        "The first practical requirement is certification. Builders will need to prove that authority artifacts are valid, invariants compile correctly, Commit Gates refuse inadmissible acts, physical gates contain consequence, and ledgers preserve enough evidence for review. The second is insurability. Risk markets will demand lineage, bounded consequence, and reconstruction before they trust autonomous infrastructure at scale.",
        "The third is domain translation. Telecommunications, healthcare, energy, finance, emergency response, logistics, public administration, and defense will not share one generic deployment pattern. Each domain has its own consequence structure and therefore its own admissibility model.",
        "The fourth is institutional literacy. Leaders must learn to distinguish useful automation from delegated authority. They must know when a system is only assisting work and when it is beginning to make consequence reachable. That is where governance has to arrive early.",
        "The governed future is not anti-autonomy. It is the condition under which autonomy can become durable. Powerful systems will scale where institutions can trust their boundaries, understand their evidence, revoke their authority, and explain their actions to the people affected by them.",
    ],
    "The Last Warrant": [
        "Power must show its Warrant. That is the final demand of the Governance Plane. Not because machines are enemies of civilization, and not because autonomy should be stopped, but because consequence without legitimacy is power without constitutional form.",
        "The Warrant is the architecture compressed into one artifact. It carries the authority chain to the point where action becomes real. It says that this act, inside this Ward, under this envelope, in this domain, during this window, under these conditions, may proceed. Without that proof, capability remains only capability. It is not legitimate power.",
        "Civilization has always required something like this. A search warrant limits state intrusion before it occurs. A judicial order binds coercive action to lawful authority. A military command carries force through a chain of command. A financial authorization binds transfer to institutional legitimacy. The autonomous age needs a machine-speed equivalent because machine action can cross into consequence before ordinary review can arrive.",
        "The Last Warrant also clarifies the human role. Human authority does not survive by requiring a person to click every approval at machine speed. It survives by embedding constitutional authorship into the runtime conditions of action. Machines may carry delegated authority, but they may not become the source of their own rightful power.",
        "Without Warrants, autonomous infrastructure tends toward silent authority accumulation. Permissions persist, exceptions normalize, agents compose agents, federation blurs responsibility, and convenience hardens into practical sovereignty. Formal institutions may remain in place while the real boundary of power moves into systems that determine what consequence becomes reachable.",
        "With Warrants, power is located. A later reviewer can ask: where did authority originate, what was the scope, which Ward applied, what did the Commit Gate decide, what evidence survived, and why was this consequence admissible now? The architecture does not promise a world without failure. It promises a world in which consequential machine action can remain bounded, reviewable, interruptible, and institutionally intelligible.",
        "That is the answer to the question at the center of the book. What must be true before an autonomous system may act? It must carry legitimate, bounded, current, action-specific authority into the boundary of consequence. At that line, power must show its Warrant.",
    ],
    "Closing Perspective": [
        "Autonomous systems force a change in where governance must live. For centuries, institutions governed infrastructure through policy, supervision, procedure, licensing, audit, and after-the-fact accountability. Those tools remain necessary, but they were built for environments where human judgment could usually stay upstream of consequence.",
        "That assumption is failing. As systems begin to sense, plan, coordinate, and act at machine speed, governance must reach the execution boundary itself. The Governance Plane is one architecture for doing so. Its claim is simple: legitimate authority must remain continuous from constitutional origin to irreversible consequence.",
        "That continuity is carried by a small set of primitives. Meta Authority Envelopes establish root legitimacy. Wards preserve protected sovereign context. Authority Domains bind action to local consequence environments. Authority Envelopes define delegated scope. Governance Invariants make constraints executable. Runtime Registers hold active state. Warrants carry action-specific authority. Commit Gates enforce admissibility. Physical Invariant Gaters protect hard boundaries. Governance Evidence Ledgers preserve institutional memory.",
        "Together, these primitives turn governance from aspiration into infrastructure. They do not make autonomy harmless. They make it governable. That is the work ahead.",
    ],
    "Model Lineage and Decision Legitimacy": [
        "A consequential autonomous act is never produced by authority alone. It is produced by a system: a model, planner, agent, controller, optimizer, or chain of cooperating components. Once those systems influence infrastructure, the institution must be able to identify not only who authorized the act, but what kind of machine participated in producing it.",
        "Model lineage is the record of that participation. It identifies the model or agent, its version, deployment source, certification state, permitted domain, trust status, and any restrictions attached to its use. This does not require every internal computation to become public. It does require the participating system to be legitimate inside the Ward where consequence occurs.",
        "The Model Lineage Certificate makes that legitimacy portable. A hospital may require medical-grade coordination agents. A telecommunications domain may require trusted routing models. An emergency Ward may permit only certified response systems. The certificate lets a Commit Gate ask whether this participant is allowed to shape this consequence.",
        "Model lineage therefore serves attribution. If a later reviewer asks why an action occurred, the answer cannot stop with the operator or the Warrant. It must also show which machine actor contributed to the decision and whether that actor was authorized to participate. Governance must reach the participants in decision formation, not only the final act.",
    ],
    "Insurable Autonomy and Institutional Trust": [
        "Insurance is one of civilization's quiet tests of trust. A system becomes insurable when risk can be bounded, attributed, reconstructed, and priced. Autonomous infrastructure will face the same test. Capability will draw interest, but insurability will decide where consequence-bearing systems can actually be deployed.",
        "The Governance Plane makes autonomy more insurable by turning consequence into an evidentiary event. The record can show authority lineage, Ward context, Warrant status, Commit Gate decision, telemetry state, physical containment, model lineage, and execution outcome. That evidence gives insurers and institutions a way to distinguish governed failure from inadmissible action.",
        "This is different from ordinary confidence in a vendor. A trusted organization may still deploy an ungovernable system. A capable model may still create unbounded exposure. Insurable autonomy requires proof that the system can narrow authority, refuse action, preserve records, and survive degraded conditions without expanding its own power.",
        "The market will eventually reward this discipline. Public agencies, infrastructure operators, hospitals, utilities, carriers, financial institutions, and defense-adjacent systems will need more than performance claims. They will need systems whose risk can be explained to boards, regulators, courts, insurers, and the public.",
    ],
    "Deterministic Actuator Enforcement": [
        "The architecture ultimately has to reach the actuator. A governance theory that stops at policy, identity, logging, or model behavior is incomplete once autonomous systems move aircraft, route power, control machinery, unlock doors, reposition vehicles, settle funds, or change the state of critical infrastructure.",
        "Deterministic Actuator Enforcement is the rule that certain commands must be refused at the final execution boundary, regardless of what the planner, model, or orchestration layer proposes. At that level, governance becomes less like advice and more like a breaker, interlock, governor, or hard stop.",
        "This does not collapse governance into safety engineering. Authority and safety remain different. A command may be physically safe and constitutionally inadmissible. Another may be authorized in principle but physically unsafe under current conditions. The architecture needs both judgments before consequence occurs.",
        "The actuator layer is where software confidence meets physical fact. If telemetry is corrupted, communications are degraded, a node is compromised, or a model behaves outside expectation, the system must still preserve non-negotiable limits. Deterministic enforcement is the last chance to keep machine power from becoming irreversible harm.",
    ],
    "Physical Invariant Gating and Consequence Containment": [
        "Physical Invariant Gating gives deterministic actuator enforcement its content. It defines the states the system must not reach: excessive torque, unsafe speed, thermal overload, pressure violation, geofence breach, collision envelope, current limit, altitude ceiling, exclusion zone, or emergency shutdown boundary.",
        "These invariants are not suggestions to a model. They are hard constraints near the place where physical consequence occurs. The system may optimize within them, but it may not reason its way around them. When an invariant is violated, the proper answer is refusal, isolation, narrowing, or hard stop.",
        "This is familiar in older infrastructure. Circuit breakers, pressure relief valves, mechanical governors, lockouts, and emergency stops all recognize that some forms of control must survive bad judgment upstream. The G-Plane extends that lesson into autonomous infrastructure by making physical containment part of governance rather than a separate afterthought.",
        "The Physical Invariant Gater therefore protects the boundary between legitimate machine action and uncontrolled physical consequence. It does not replace Wards, Warrants, or Commit Gates. It gives them a hard floor: even valid authority cannot authorize a system to violate the physical conditions required for survivable operation.",
    ],
    "Emergency Governance Modes and Constitutional Survival": [
        "Emergencies are where governance is most tempted to disappear. Communications fail, telemetry fragments, operators improvise, priorities change, and urgency argues for broader authority. History shows the danger: emergency power often expands faster than accountability can follow.",
        "Emergency Governance Modes exist to prevent crisis from becoming a loophole. They allow authority to adapt under stress, but only through bounded pathways. A system may narrow action windows, elevate escalation requirements, shift into emergency posture, preserve local safe operation, or suspend high-consequence acts until authority is clear.",
        "The central rule is that emergency conditions may change how authority is exercised, but they may not erase the need for authority. A wildfire, outage, attack, or medical crisis may justify different admissibility thresholds. It does not justify invisible sovereignty.",
        "This is constitutional survival in operational form. The architecture must keep institutions alive inside crisis, even when direct supervision is degraded. Emergency mode should make authority more explicit, not less.",
    ],
    "Human Escalation, Override, and Institutional Continuity": [
        "Human authority cannot mean constant human clicking. At machine speed, that fantasy fails. But the opposite fantasy is just as dangerous: that removing humans entirely makes governance stronger. Institutions remain the source of legitimacy, and the architecture must preserve their power to interrupt, narrow, suspend, and recover control.",
        "Human escalation is therefore not a manual override button bolted onto an autonomous system. It is a set of protected pathways through which legitimate institutional authority can re-enter runtime operation when conditions change. Those pathways must be bounded, evidenced, and resistant to being bypassed by convenience.",
        "The system should know when escalation is required: degraded telemetry, conflicting authority, emergency posture, physical risk, model uncertainty, revocation ambiguity, or consequence beyond the active envelope. In those moments, autonomy should narrow rather than improvise sovereignty.",
        "Institutional continuity survives when human authority remains structurally reachable even if humans cannot supervise every act. The point is not to slow all execution to human tempo. The point is to ensure that machine tempo never severs the institution's right to say stop.",
    ],
    "The Claims of the Architecture": [
        "The architecture makes five claims. First, autonomous consequence requires present-tense legitimacy. Second, legitimacy must be carried by artifacts the machine can evaluate. Third, authority must narrow as it approaches execution. Fourth, physical limits must survive software confidence. Fifth, evidence must make the act reconstructable after the fact.",
        "These claims are not separate theories. They are the same thesis viewed from different points in the stack. Wards protect context. Warrants prove action-level authority. Commit Gates enforce admissibility. Physical Invariant Gaters preserve containment. Evidence Ledgers preserve institutional memory.",
        "The claim of the G-Plane is therefore not that machines should be weak. It is that powerful machines must remain governable. Capability is not sovereignty. Usefulness is not legitimacy. Execution must show its authority.",
    ],
    "The Execution Constitution": [
        "The execution constitution is the portion of governance that reaches the boundary of action. It is not the whole institutional order. It is the subset of authority, constraint, state, revocation, and evidence that must be present when a consequential system is about to act.",
        "This matters because many institutions possess rich governance language that never becomes operational. The G-Plane asks what part of that language can be compiled into runtime conditions and tested before consequence.",
        "The execution constitution is therefore practical: active root authority, protected domain, delegated scope, invariant set, runtime state, Warrant, Commit Gate decision, physical containment, and evidence obligation. If those elements are absent, governance has not reached execution.",
    ],
    "The Governance Gap": [
        "The governance gap is the distance between machine-speed consequence and institution-speed authority. Modern infrastructure can reroute traffic, allocate compute, settle transactions, isolate networks, move robots, or alter access before ordinary institutional review can occur.",
        "Monitoring does not close that gap. Monitoring tells an institution what happened. Governance determines what may happen. Once consequence has crossed into infrastructure, the institution may investigate, punish, compensate, or redesign, but it has already lost the upstream boundary.",
        "The gap is therefore temporal before it is technical. Autonomous systems compress the time between decision and consequence. The G-Plane exists to put legitimacy back into that compressed interval.",
    ],
    "Architecture Instead of Policy": [
        "Policy is necessary, but policy alone does not bind a machine at the moment of action. A written rule becomes operational only when it can be represented, evaluated, enforced, and recorded inside the system that reaches consequence.",
        "Architecture is the translation layer. It turns institutional authority into artifacts, artifacts into constraints, constraints into runtime state, runtime state into admissibility decisions, and admissibility decisions into evidence.",
        "The move from policy to architecture is not anti-law or anti-institution. It is the opposite. It is the attempt to make institutional authority survive inside systems that act too quickly for paper governance to remain upstream.",
    ],
    "Governance Invariant Compilation": [
        "Governance Invariants are the machine-checkable constraints produced from authority. They are not policy language. They are the executable conditions that must remain true before an action may become admissible.",
        "Compilation matters because institutional rules usually arrive in human form: charters, policies, safety limits, regulations, contracts, operational orders, domain rules, and emergency procedures. The G-Plane has to translate those sources into deterministic conditions without pretending the translation is the source of legitimacy.",
        "An invariant is therefore a compiled expression of authority, not authority itself. Its job is narrow and essential: make the relevant rule testable at machine speed.",
    ],
    "Runtime Governance Registers": [
        "Runtime Registers hold the live state against which invariants are evaluated. If invariants are the compiled conditions, registers are the present-tense facts: active authority, current telemetry, revocation exposure, synchronization status, Warrant state, emergency mode, and physical-gate status.",
        "This distinction separates rules from state. A rule may be valid while the current state makes execution inadmissible. A Warrant may exist while telemetry or revocation conditions make it unusable. The register is where that present-tense reality becomes available to the Commit Gate.",
        "Registers are therefore not another policy artifact. They are the live memory of the governance system at the edge of action.",
    ],
    "Multi-Agent Governance and Shared Consequence Coordination": [
        "Multi-agent governance addresses a problem that single-system governance cannot solve: several autonomous systems may contribute to one consequence. No individual agent may appear to hold the whole decision, yet the combined action may alter infrastructure state.",
        "Shared consequence requires shared admissibility. Each participating system must carry compatible authority, and the coordinated act must be evaluated as a whole. Otherwise responsibility fragments across agents, models, services, operators, and institutions.",
        "The G-Plane treats coordination as a governance event, not merely an orchestration event. The question is not only whether agents can cooperate. The question is whether their combined action remains legitimate.",
    ],
    "Governance Plane Topology and Distributed Runtime Architecture": [
        "Topology describes where governance functions live. Some authority remains centralized. Some state must be local. Some evidence can be federated. Some physical gates sit at the actuator. Distributed runtime architecture decides how those pieces are placed.",
        "The topology cannot assume one perfect center. Infrastructure is too distributed, too latency-sensitive, and too jurisdictionally fragmented. The architecture must decide which functions require root authority, which can be cached locally, which must synchronize, and which must fail closed.",
        "This chapter is about placement: root services, domain registries, local kernels, edge Commit Gates, physical interlocks, witness nodes, and evidence anchors. A governance architecture that cannot be placed cannot be deployed.",
    ],
    "Governance Witness Systems and Distributed Legitimacy Consensus": [
        "Witness systems answer a different question from topology. Topology asks where governance functions live. Witnessing asks who can attest that governance remained valid when distributed systems acted.",
        "A witness may attest that a Warrant existed, that revocation had not reached a node, that a Commit Gate decision matched active state, that a ledger record was anchored, or that a domain operated under degraded but bounded authority. Witnessing gives distributed governance an evidentiary surface outside the acting system itself.",
        "The point is not to create perfect consensus over every action. The point is to create enough independent legitimacy evidence that a consequential act can be reconstructed, challenged, insured, or refused across domains.",
    ],
}


def has_image(paragraph: Paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
            return True
    return False


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    ppr = paragraph._p.pPr
    for child in list(paragraph._p):
        if ppr is not None and child is ppr:
            continue
        paragraph._p.remove(child)
    paragraph.add_run(text)


def replace_chapter_body(doc: Document, chapter_title: str, new_paragraphs: list[str]) -> bool:
    paragraphs = doc.paragraphs
    title = None
    for paragraph in paragraphs:
        if paragraph.text.strip() == chapter_title and (paragraph.style.name if paragraph.style else "").startswith("Heading 1"):
            title = paragraph
            break
    if title is None:
        return False

    body = []
    images = []
    capture = False
    for paragraph in list(doc.paragraphs):
        if paragraph._element is title._element:
            capture = True
            continue
        if not capture:
            continue
        style = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        if style.startswith("Heading") and text:
            break
        if has_image(paragraph) or text.startswith("Figure 0."):
            images.append(paragraph)
        elif text:
            body.append(paragraph)

    anchor = title
    inserted = []
    for idx, text in enumerate(new_paragraphs):
        anchor = insert_after(anchor, text, "First Paragraph" if idx == 0 else "Body Text")
        inserted.append(anchor)

    # Keep chapter figures after the second paragraph where possible.
    if images:
        figure_anchor = inserted[min(1, len(inserted) - 1)]
        for image_para in images:
            image_para._p.getparent().remove(image_para._p)
            figure_anchor._p.addnext(image_para._p)
            figure_anchor = image_para

    for paragraph in body:
        if paragraph._element is not None:
            delete_paragraph(paragraph)
    return True


def remove_stock_repetition(doc: Document) -> int:
    patterns = [
        "attempts to operationalize these requirements directly",
        "may ultimately become one of the defining",
        "The distinction may ultimately become one of",
        "The architecture attempts to operationalize these requirements",
        "Civilization-scale autonomous infrastructure systems increasingly require governance architectures capable",
        "This layered continuity",
        "difference carries profound consequence",
        "important once autonomous begin",
    ]
    removed = 0
    for paragraph in list(doc.paragraphs):
        if paragraph._element is None or has_image(paragraph):
            continue
        text = paragraph.text.strip()
        if len(text) < 900 and any(pattern in text for pattern in patterns):
            delete_paragraph(paragraph)
            removed += 1
    return removed


def clean_forbidden_phrasing(doc: Document) -> int:
    replacements = {
        "That distinction matters.": "That distinction is central.",
        "Compilation matters because": "Compilation is necessary because",
        "This distinction matters because": "The distinction is important because",
        "This matters because": "The reason is simple:",
        "the boundary where control matters": "the boundary where control is real",
        "Local control mattered, common coordination also mattered, and security mattered. Commerce mattered, jurisdiction mattered, and legitimacy mattered.": "Local control, common coordination, security, commerce, jurisdiction, and legitimacy all had to coexist.",
        "aautonomous": "autonomous",
        "Aautonomous": "Autonomous",
        "aautomation": "automation",
        "Aautomation": "Automation",
        "aautonomy": "autonomy",
        "Aautonomy": "Autonomy",
        "ccentralized": "centralized",
        "Ccentralized": "Centralized",
        "ccomputational": "computational",
        "Ccomputational": "Computational",
        "ffederated": "federated",
        "Ffederated": "Federated",
        "ggovernable": "governable",
        "Ggovernable": "Governable",
        "jjurisdictional": "jurisdictional",
        "Jjurisdictional": "Jurisdictional",
        "mmmmachine": "machine",
        "Mmmmachine": "Machine",
        "mmmachine": "machine",
        "Mmmachine": "Machine",
        "mmachine": "machine",
        "Mmachine": "Machine",
        "cccomputational": "computational",
        "Cccomputational": "Computational",
        "ttransnational": "transnational",
        "Ttransnational": "Transnational",
    }
    changed = 0
    for paragraph in doc.paragraphs:
        if paragraph._element is None or has_image(paragraph):
            continue
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(paragraph, new_text)
            changed += 1
    return changed


def remove_duplicate_plain_paragraphs(doc: Document) -> int:
    seen = set()
    removed = 0
    for paragraph in list(doc.paragraphs):
        if paragraph._element is None or has_image(paragraph):
            continue
        text = paragraph.text.strip()
        if not text or paragraph.style.name.startswith("Heading"):
            continue
        key = re.sub(r"\s+", " ", text)
        if key in seen:
            delete_paragraph(paragraph)
            removed += 1
        else:
            seen.add(key)
    return removed


doc = Document(str(DOCX))
changed = 0
for title, paragraphs in CHAPTER_REWRITES.items():
    if replace_chapter_body(doc, title, paragraphs):
        changed += 1
removed = remove_stock_repetition(doc)
phrasing = clean_forbidden_phrasing(doc)
deduped = remove_duplicate_plain_paragraphs(doc)

doc.core_properties.comments = "Repetition reduction pass: differentiated duplicated chapter jobs and removed stock repeated closures."
doc.save(str(DOCX))
print(f"chapters_rewritten={changed} paragraphs_removed={removed} phrasing_cleanups={phrasing} duplicate_paragraphs_removed={deduped} output={DOCX}")
